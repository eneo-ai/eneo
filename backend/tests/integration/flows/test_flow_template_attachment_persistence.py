from __future__ import annotations

import io
from collections.abc import AsyncIterator
from uuid import uuid4

import pytest
import sqlalchemy as sa
from docx import Document
from sqlalchemy.exc import DBAPIError, IntegrityError
from sqlalchemy.orm.session import SessionTransaction

from eneo.audit.domain.action_types import ActionType
from eneo.audit.domain.audit_log import AuditLog
from eneo.audit.domain.entity_types import EntityType
from eneo.audit.infrastructure.audit_log_repo_impl import AuditLogRepositoryImpl
from eneo.database.database import sessionmanager
from eneo.database.tables.audit_log_table import AuditLog as AuditLogTable
from eneo.database.tables.files_table import Files
from eneo.database.tables.flow_tables import (
    BuilderSessionFiles,
    BuilderSessions,
    Flows,
    FlowTemplateAssets,
    FlowVersions,
)
from eneo.database.tables.spaces_table import Spaces
from eneo.files.file_models import FileContentVariant, FileType
from eneo.files.file_protocol import PendingFileContent, PreparedFileUpload
from eneo.flows.ai_builder.ai_builder_domain_models import TargetKind
from eneo.flows.ai_builder.ai_builder_repo import AIBuilderRepository
from eneo.flows.application.flow_authoring_command import TemplateAttachmentIntent
from eneo.flows.application.flow_draft_materialization import (
    FlowDraftChangeSet,
    FlowDraftCompiledStep,
    FlowDraftStepChangeKind,
    compile_flow_draft_changeset,
)
from eneo.flows.application.flow_template_attachment_materialization import (
    materialize_template_attachment,
)
from eneo.flows.domain.flow import FlowStep
from eneo.flows.flow_api_error_code import FlowApiErrorCode
from eneo.flows.flow_authoring_spec import (
    AssistantSpec,
    FlowDraftSpecCore,
    InputSource,
    InputType,
    OutputMode,
    OutputType,
    StepSpec,
)
from eneo.flows.flow_resource_bindings import FlowResourceBindingSource
from eneo.flows.flow_template_asset_service import (
    AttachedTemplateFileUnavailableError,
)
from eneo.main.exceptions import BadRequestException

DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


def _template_bytes(placeholder: str) -> bytes:
    document = Document()
    document.add_paragraph("{{ " + placeholder + " }}")
    payload = io.BytesIO()
    document.save(payload)
    return payload.getvalue()


async def _bytes(payload: bytes) -> AsyncIterator[bytes]:
    yield payload


async def _save_template_file(container, *, name: str, content: bytes):
    return await container.file_service().save_prepared_file(
        PreparedFileUpload(
            name=name,
            file_type=FileType.DOCUMENT,
            display_media_type=DOCX_MIME,
            contents=(
                PendingFileContent(
                    variant=FileContentVariant.ORIGINAL,
                    chunks=_bytes(content),
                    declared_media_type=DOCX_MIME,
                    verified_media_type=DOCX_MIME,
                ),
            ),
        )
    )


def _template_changeset():
    spec = FlowDraftSpecCore(
        flow_name="Template flow",
        steps=[
            StepSpec(
                plan_step_ref="step_a",
                name="Prepare",
                assistant_spec=AssistantSpec(instructions="Prepare content."),
                input_source=InputSource.FLOW_INPUT,
                input_type=InputType.TEXT,
            ),
            StepSpec(
                plan_step_ref="step_b",
                name="Fill",
                assistant_spec=AssistantSpec(instructions="Fill template."),
                input_source=InputSource.PREVIOUS_STEP,
                input_type=InputType.TEXT,
                output_mode=OutputMode.TEMPLATE_FILL,
                output_type=OutputType.DOCX,
                output_config={"bindings": {"case_id": "{{ flow_input.case_id }}"}},
            ),
        ],
    )
    return compile_flow_draft_changeset(spec, current_flow=None)


async def _create_flow_and_file(container, *, placeholder: str):
    user = container.user()
    space = Spaces(
        tenant_id=user.tenant_id,
        user_id=user.id,
        name=f"template-binding-{uuid4().hex}",
    )
    container.session().add(space)
    await container.session().flush()
    flow = await container.flow_service().create_flow(
        space_id=space.id,
        name=f"Template flow {uuid4().hex}",
        description="",
        steps=[],
    )
    content = _template_bytes(placeholder)
    file = await _save_template_file(
        container,
        name="template.docx",
        content=content,
    )
    return user, space, flow, file


async def _create_template_download_fixture(db_container):
    async with db_container() as setup_container:
        user, _space, flow, file = await _create_flow_and_file(
            setup_container,
            placeholder="case_id",
        )
        token = setup_container.auth_service().create_access_token_for_user(user)

    async with db_container(user=user) as asset_container:
        asset = await asset_container.flow_template_asset_service().create_from_existing_attached_file(
            flow_id=flow.id,
            file_id=file.id,
        )
    return token, flow.id, asset.id, file.id, user.id, user.tenant_id


async def _delete_file(
    *,
    file_id,
    lock_timeout: bool,
) -> None:
    async with sessionmanager.session() as session, session.begin():
        if lock_timeout:
            await session.execute(sa.text("SET LOCAL lock_timeout = '100ms'"))
        await session.execute(sa.delete(Files).where(Files.id == file_id))


@pytest.mark.integration
@pytest.mark.asyncio
async def test_template_download_url_commits_bounded_audit_before_response(
    client,
    db_container,
    patch_auth_service_jwt,
) -> None:
    _ = patch_auth_service_jwt
    (
        token,
        flow_id,
        asset_id,
        file_id,
        user_id,
        tenant_id,
    ) = await _create_template_download_fixture(db_container)

    response = await client.post(
        f"/api/v1/flows/{flow_id}/template-files/{asset_id}/signed-url/",
        json={"expires_in": 120},
        headers={"Authorization": f"Bearer {token}"},
    )

    assert response.status_code == 200, response.text
    async with db_container() as container:
        audit_log = await container.session().scalar(
            sa.select(AuditLogTable).where(
                AuditLogTable.action == ActionType.FILE_SIGNED_URL_MINTED.value,
                AuditLogTable.entity_type == EntityType.FILE.value,
                AuditLogTable.entity_id == file_id,
            )
        )
        assert audit_log is not None
        assert audit_log.tenant_id == tenant_id
        assert audit_log.actor_id == user_id
        assert audit_log.actor_api_key_id is None
        assert audit_log.actor_type == "user"
        assert audit_log.outcome == "success"
        assert audit_log.log_metadata["extra"] == {
            "flow_id": str(flow_id),
            "template_asset_id": str(asset_id),
            "file_id": str(file_id),
            "download_purpose": "flow_template",
        }
        assert response.json()["url"] not in repr(audit_log.log_metadata)


@pytest.mark.integration
@pytest.mark.asyncio
async def test_template_download_audit_commit_failure_returns_typed_503(
    client,
    db_container,
    patch_auth_service_jwt,
    monkeypatch,
) -> None:
    _ = patch_auth_service_jwt
    (
        token,
        flow_id,
        asset_id,
        file_id,
        _user_id,
        _tenant_id,
    ) = await _create_template_download_fixture(db_container)
    events: list[str] = []
    original_create_audit = AuditLogRepositoryImpl.create
    original_commit = SessionTransaction.commit

    async def mark_template_download_audit_transaction(
        repository: AuditLogRepositoryImpl,
        audit_log: AuditLog,
    ) -> AuditLog:
        result = await original_create_audit(repository, audit_log)
        if audit_log.action == ActionType.FILE_SIGNED_URL_MINTED:
            repository.session.sync_session.info["fail_template_audit_commit"] = True
        return result

    def fail_template_download_audit_commit(
        transaction: SessionTransaction,
        *,
        _to_root: bool = False,
    ) -> None:
        if transaction.session.info.pop("fail_template_audit_commit", False):
            events.append("commit_failed")
            raise RuntimeError("template download audit commit unavailable")
        original_commit(transaction, _to_root=_to_root)

    monkeypatch.setattr(
        AuditLogRepositoryImpl,
        "create",
        mark_template_download_audit_transaction,
    )
    monkeypatch.setattr(
        SessionTransaction, "commit", fail_template_download_audit_commit
    )

    response = await client.post(
        f"/api/v1/flows/{flow_id}/template-files/{asset_id}/signed-url/",
        json={"expires_in": 120},
        headers={"Authorization": f"Bearer {token}"},
    )

    assert response.status_code == 503
    assert events == ["commit_failed"]
    payload = response.json()
    assert payload["code"] == FlowApiErrorCode.TEMPLATE_DOWNLOAD_AUDIT_UNAVAILABLE
    assert payload["context"] == {"audit_required": True}
    assert "url" not in payload
    async with db_container() as container:
        audit_count = await container.session().scalar(
            sa.select(sa.func.count(AuditLogTable.id)).where(
                AuditLogTable.action == ActionType.FILE_SIGNED_URL_MINTED.value,
                AuditLogTable.entity_id == file_id,
            )
        )
    assert audit_count == 0


@pytest.mark.integration
@pytest.mark.asyncio
async def test_template_download_hides_cross_tenant_asset_without_audit(
    client,
    db_container,
    patch_auth_service_jwt,
    tenant_factory,
    user_factory,
) -> None:
    _ = patch_auth_service_jwt
    async with db_container() as setup_container:
        current_user = setup_container.user()
        token = setup_container.auth_service().create_access_token_for_user(
            current_user
        )
        other_tenant = await tenant_factory(
            setup_container.session(),
            name=f"Flow audit isolation {uuid4().hex}",
        )
        other_user = await user_factory(
            setup_container.session(),
            tenant_id=other_tenant.id,
        )
        other_space = Spaces(
            tenant_id=other_tenant.id,
            user_id=other_user.id,
            name=f"Other tenant Flow space {uuid4().hex}",
        )
        setup_container.session().add(other_space)
        await setup_container.session().flush()
        other_flow = Flows(
            tenant_id=other_tenant.id,
            space_id=other_space.id,
            name=f"Other tenant Flow {uuid4().hex}",
            description="Tenant-isolation audit fixture.",
            created_by_user_id=other_user.id,
            owner_user_id=other_user.id,
        )
        other_file = Files(
            tenant_id=other_tenant.id,
            owner_type="user",
            owner_user_id=other_user.id,
            name="other-tenant-template.docx",
            mimetype=DOCX_MIME,
            file_type=FileType.DOCUMENT.value,
        )
        setup_container.session().add_all([other_flow, other_file])
        await setup_container.session().flush()
        other_asset = FlowTemplateAssets(
            tenant_id=other_tenant.id,
            space_id=other_space.id,
            flow_id=other_flow.id,
            file_id=other_file.id,
            name=other_file.name,
            checksum="cross-tenant-template-fixture",
            mimetype=DOCX_MIME,
            placeholders=[],
            created_by_user_id=other_user.id,
            updated_by_user_id=other_user.id,
        )
        setup_container.session().add(other_asset)
        await setup_container.session().flush()
        other_flow_id = other_flow.id
        other_asset_id = other_asset.id
        other_file_id = other_file.id

    response = await client.post(
        (f"/api/v1/flows/{other_flow_id}/template-files/{other_asset_id}/signed-url/"),
        json={"expires_in": 120},
        headers={"Authorization": f"Bearer {token}"},
    )

    assert response.status_code == 404
    assert "url" not in response.json()
    async with db_container() as container:
        audit_count = await container.session().scalar(
            sa.select(sa.func.count(AuditLogTable.id)).where(
                AuditLogTable.action == ActionType.FILE_SIGNED_URL_MINTED.value,
                AuditLogTable.entity_id == other_file_id,
            )
        )
    assert audit_count == 0


@pytest.mark.integration
@pytest.mark.asyncio
async def test_promoted_template_survives_builder_session_deletion_and_fences_file(
    db_container,
) -> None:
    async with db_container() as setup_container:
        user, space, flow, file = await _create_flow_and_file(
            setup_container,
            placeholder="case_id",
        )
        repo = AIBuilderRepository(setup_container.session())
        builder_session = await repo.create_session(
            tenant_id=user.tenant_id,
            space_id=space.id,
            actor_user_id=user.id,
            target_kind=TargetKind.CREATE,
        )
        await setup_container.session().execute(
            sa.insert(BuilderSessionFiles).values(
                session_id=builder_session.id,
                file_id=file.id,
                tenant_id=user.tenant_id,
            )
        )

    async with db_container(user=user) as container:
        asset = await container.flow_template_asset_service().create_from_existing_attached_file(
            flow_id=flow.id,
            file_id=file.id,
        )
        await container.session().execute(
            sa.delete(BuilderSessions).where(BuilderSessions.id == builder_session.id)
        )
        await container.session().flush()

        persisted_asset_file_id = await container.session().scalar(
            sa.select(FlowTemplateAssets.file_id).where(
                FlowTemplateAssets.id == asset.id
            )
        )
        persisted_file_id = await container.session().scalar(
            sa.select(Files.id).where(Files.id == file.id)
        )
        assert persisted_asset_file_id == file.id
        assert persisted_file_id == file.id

        with pytest.raises(IntegrityError):
            async with container.session().begin_nested():
                await container.session().execute(
                    sa.delete(Files).where(Files.id == file.id)
                )
                await container.session().flush()


@pytest.mark.integration
@pytest.mark.asyncio
async def test_template_promotion_lock_serializes_both_file_delete_race_orders(
    db_container,
) -> None:
    async with db_container() as setup_container:
        user, _space, flow, file = await _create_flow_and_file(
            setup_container,
            placeholder="case_id",
        )

    async with db_container(user=user) as promotion_container:
        asset = await promotion_container.flow_template_asset_service().create_from_existing_attached_file(
            flow_id=flow.id,
            file_id=file.id,
        )
        with pytest.raises(DBAPIError) as lock_error:
            await _delete_file(file_id=file.id, lock_timeout=True)
        assert getattr(lock_error.value.orig, "sqlstate", None) == "55P03"
        assert asset.file_id == file.id

    with pytest.raises(IntegrityError):
        await _delete_file(file_id=file.id, lock_timeout=False)

    async with db_container(user=user) as setup_container:
        second_content = _template_bytes("reference_number")
        second_file = await _save_template_file(
            setup_container,
            name="second-template.docx",
            content=second_content,
        )
    await _delete_file(file_id=second_file.id, lock_timeout=False)

    async with db_container(user=user) as promotion_container:
        with pytest.raises(AttachedTemplateFileUnavailableError):
            await promotion_container.flow_template_asset_service().create_from_existing_attached_file(
                flow_id=flow.id,
                file_id=second_file.id,
            )


@pytest.mark.integration
@pytest.mark.asyncio
async def test_changed_template_contract_rolls_back_promoted_asset(
    db_container,
) -> None:
    async with db_container() as setup_container:
        user, _space, flow, file = await _create_flow_and_file(
            setup_container,
            placeholder="customer.name",
        )

    async with db_container(user=user) as container:
        with pytest.raises(BadRequestException) as exc_info:
            async with container.session().begin_nested():
                await materialize_template_attachment(
                    intent=TemplateAttachmentIntent(
                        file_id=file.id,
                        terminal_plan_step_ref="step_b",
                    ),
                    changeset=_template_changeset(),
                    flow_id=flow.id,
                    template_asset_service=container.flow_template_asset_service(),
                )

        assert exc_info.value.code == "architecture_materialization_failed"
        asset_count = await container.session().scalar(
            sa.select(sa.func.count(FlowTemplateAssets.id)).where(
                FlowTemplateAssets.flow_id == flow.id,
                FlowTemplateAssets.file_id == file.id,
            )
        )
        assert asset_count == 0
        assert (
            await container.session().scalar(
                sa.select(Files.id).where(Files.id == file.id)
            )
            == file.id
        )


@pytest.mark.integration
@pytest.mark.asyncio
async def test_materialized_template_attachment_publishes_with_pinned_identity(
    db_container,
) -> None:
    async with db_container() as setup_container:
        user, _space, flow, file = await _create_flow_and_file(
            setup_container,
            placeholder="case_id",
        )

    async with db_container(user=user) as container:
        changeset = FlowDraftChangeSet(
            flow_name=flow.name,
            flow_description="",
            compiled_steps=[
                FlowDraftCompiledStep(
                    plan_step_ref="step_a",
                    change_kind=FlowDraftStepChangeKind.ADDED,
                    step_order=1,
                    user_description="Fill template",
                    input_source="flow_input",
                    input_type="text",
                    output_mode="template_fill",
                    output_type="docx",
                    output_config={"bindings": {"case_id": "{{ flow_input.case_id }}"}},
                )
            ],
        )
        materialized = await materialize_template_attachment(
            intent=TemplateAttachmentIntent(
                file_id=file.id,
                terminal_plan_step_ref="step_a",
            ),
            changeset=changeset,
            flow_id=flow.id,
            template_asset_service=container.flow_template_asset_service(),
        )
        assistant, _ = await container.flow_service().create_flow_assistant(
            flow_id=flow.id,
            name="template-fill",
        )
        terminal = materialized.changeset.compiled_steps[0]
        updated = await container.flow_service().update_flow(
            flow_id=flow.id,
            metadata_json=materialized.changeset.metadata_json,
            steps=[
                FlowStep(
                    assistant_id=assistant.id,
                    step_order=terminal.step_order,
                    user_description=terminal.user_description,
                    input_source=terminal.input_source,
                    input_type=terminal.input_type,
                    output_mode=terminal.output_mode,
                    output_type=terminal.output_type,
                    output_config=terminal.output_config,
                )
            ],
        )
        await container.flow_service().replace_resource_bindings(
            flow_id=flow.id,
            bindings=(materialized.binding,),
            source=FlowResourceBindingSource.AI_BUILDER,
        )

        published = await container.flow_service().publish_flow(flow_id=updated.id)

        assert published.published_version == 1
        definition = await container.session().scalar(
            sa.select(FlowVersions.definition_json).where(
                FlowVersions.flow_id == flow.id,
                FlowVersions.tenant_id == user.tenant_id,
                FlowVersions.version == 1,
            )
        )
        assert isinstance(definition, dict)
        output_config = definition["steps"][0]["output_config"]
        assert output_config["template_asset_id"] == str(materialized.binding.local_id)
        assert output_config["template_checksum"] == file.checksum
        assert output_config["placeholders"] == ["case_id"]
        assert output_config["bindings"] == {"case_id": "{{ flow_input.case_id }}"}
