from __future__ import annotations

import io
import sys
from types import SimpleNamespace
import zipfile

import pytest
from docx import Document
from docx.oxml.ns import qn

from intric.flows.runtime.docx_template_runtime import (
    _iter_section_headers,
    _iter_section_footers,
    extract_docx_template_text_preview,
    extract_docx_text,
    inspect_docx_template_bytes,
    render_docx_template,
)
from intric.main.exceptions import FileNotSupportedException, TypedIOValidationException


def _build_template_bytes() -> bytes:
    doc = Document()
    doc.add_paragraph("Titel: {{title}}")
    split = doc.add_paragraph()
    split.add_run("{{")
    split.add_run("introduction")
    split.add_run("}}")

    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "{{methodology}}"

    header = doc.sections[0].header
    header.add_paragraph("Rapport för {{author}}")

    footer = doc.sections[0].footer
    footer.add_paragraph("Version {{version}}")

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def test_inspect_docx_template_bytes_finds_placeholders_across_document_regions() -> None:
    placeholders = inspect_docx_template_bytes(_build_template_bytes(), filename="rapport.docx")

    assert {(item["name"], item["location"]) for item in placeholders} >= {
        ("title", "body"),
        ("introduction", "body"),
        ("methodology", "table"),
        ("author", "header"),
        ("version", "footer"),
    }


def test_render_docx_template_replaces_split_run_placeholders() -> None:
    blob, mimetype, filename = render_docx_template(
        template_bytes=_build_template_bytes(),
        context={
            "title": "Demokratiskt deltagande",
            "introduction": "Inledningstext",
            "methodology": "Metodavsnitt",
            "author": "Anders Svensson",
            "version": "3",
        },
        step_order=11,
    )

    rendered = Document(io.BytesIO(blob))
    all_text = "\n".join(paragraph.text for paragraph in rendered.paragraphs)
    header_text = "\n".join(paragraph.text for paragraph in rendered.sections[0].header.paragraphs)
    footer_text = "\n".join(paragraph.text for paragraph in rendered.sections[0].footer.paragraphs)

    assert mimetype == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    assert filename == "step_11_output.docx"
    assert "Demokratiskt deltagande" in all_text
    assert "Inledningstext" in all_text
    assert rendered.tables[0].cell(0, 0).text == "Metodavsnitt"
    assert "Anders Svensson" in header_text
    assert "Version 3" in footer_text


def test_render_docx_template_supports_flat_context_for_dotted_placeholder_names() -> None:
    doc = Document()
    doc.add_paragraph("{{step_1.output.text}}")
    buffer = io.BytesIO()
    doc.save(buffer)

    blob, _, _ = render_docx_template(
        template_bytes=buffer.getvalue(),
        context={"step_1.output.text": "Dotted binding rendered"},
        step_order=12,
    )

    rendered = Document(io.BytesIO(blob))
    assert "Dotted binding rendered" in "\n".join(paragraph.text for paragraph in rendered.paragraphs)


def test_render_docx_template_supports_placeholders_with_spaces() -> None:
    doc = Document()
    doc.add_paragraph("Rubrik: {{Planering och Hälsa}}")
    buffer = io.BytesIO()
    doc.save(buffer)

    blob, _, _ = render_docx_template(
        template_bytes=buffer.getvalue(),
        context={"Planering och Hälsa": "Kort sammanfattning"},
        step_order=13,
    )

    rendered = Document(io.BytesIO(blob))
    assert "Rubrik: Kort sammanfattning" in "\n".join(
        paragraph.text for paragraph in rendered.paragraphs
    )


def test_inspect_docx_template_bytes_rejects_macro_enabled_payloads() -> None:
    template_bytes = _build_template_bytes()
    mutated = io.BytesIO(template_bytes)
    with zipfile.ZipFile(mutated, mode="a") as archive:
        archive.writestr("word/vbaProject.bin", b"macro")

    with pytest.raises(
        FileNotSupportedException,
        match="(?i)macro-enabled",
    ) as exc_info:
        inspect_docx_template_bytes(mutated.getvalue(), filename="rapport.docm")

    assert exc_info.value.code == "flow_template_macro_not_allowed"


def test_render_docx_template_rejects_missing_context_placeholders() -> None:
    with pytest.raises(TypedIOValidationException, match="Unresolved template placeholders"):
        render_docx_template(
            template_bytes=_build_template_bytes(),
            context={"title": "Only title"},
            step_order=4,
        )


def test_extract_docx_template_text_preview_returns_readable_text() -> None:
    preview = extract_docx_template_text_preview(_build_template_bytes())

    assert "Titel: {{title}}" in preview
    assert "Rapport för {{author}}" in preview


def test_render_docx_template_wraps_undefined_render_errors_with_actionable_message(monkeypatch) -> None:
    class _FakeDocxTemplate:
        def __init__(self, _path: str) -> None:
            pass

        def render(self, _context, jinja_env=None) -> None:
            raise RuntimeError("'summary' is undefined")

        def save(self, _path: str) -> None:
            raise AssertionError("save should not be reached")

    monkeypatch.setitem(sys.modules, "docxtpl", SimpleNamespace(DocxTemplate=_FakeDocxTemplate))

    with pytest.raises(TypedIOValidationException, match="written directly in the DOCX"):
        render_docx_template(
            template_bytes=_build_template_bytes(),
            context={
                "title": "Demokratiskt deltagande",
                "introduction": "Inledningstext",
                "methodology": "Metodavsnitt",
                "author": "Anders Svensson",
                "version": "3",
            },
            step_order=8,
        )


# ---------------------------------------------------------------------------
# Header/footer traversal edge cases (regression for FileNotFoundError crash)
# ---------------------------------------------------------------------------


def _build_no_header_footer_bytes() -> bytes:
    """Build a DOCX with body content only — no headers or footers."""
    doc = Document()
    doc.add_paragraph("Body text with {{placeholder}}")
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _build_first_page_header_only_bytes() -> bytes:
    """Build a DOCX with a first-page header but no default header."""
    doc = Document()
    doc.add_paragraph("Body: {{body_var}}")
    section = doc.sections[0]
    section.different_first_page_header_footer = True
    first_hdr = section.first_page_header
    first_hdr.add_paragraph("First page: {{first_var}}")
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _build_even_page_header_bytes() -> bytes:
    """Build a DOCX with both default and even-page headers."""
    doc = Document()
    doc.add_paragraph("Body: {{body_var}}")
    section = doc.sections[0]
    # Add default header
    section.header.add_paragraph("Default: {{default_hdr}}")
    # Add even-page header (requires document setting)
    even_hdr = section.even_page_header
    even_hdr.add_paragraph("Even page: {{even_hdr}}")
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def test_inspect_no_header_no_footer_does_not_crash() -> None:
    """Regression: DOCX without headers/footers must not trigger FileNotFoundError."""
    template = _build_no_header_footer_bytes()
    placeholders = inspect_docx_template_bytes(template, filename="plain.docx")

    names = {item["name"] for item in placeholders}
    assert "placeholder" in names
    # No header/footer placeholders expected
    locations = {item["location"] for item in placeholders}
    assert "header" not in locations
    assert "footer" not in locations


def test_extract_text_no_header_no_footer_does_not_crash() -> None:
    """Regression: extract_docx_text on headerless DOCX must not crash."""
    text = extract_docx_text(_build_no_header_footer_bytes())
    assert "Body text" in text


def test_inspect_finds_first_page_header_placeholders() -> None:
    """First-page headers must be inspected — not just default headers."""
    template = _build_first_page_header_only_bytes()
    placeholders = inspect_docx_template_bytes(template, filename="first.docx")

    names = {item["name"] for item in placeholders}
    assert "first_var" in names
    assert "body_var" in names


def test_inspect_finds_even_page_header_placeholders() -> None:
    """Even-page headers must be inspected alongside default headers."""
    template = _build_even_page_header_bytes()
    placeholders = inspect_docx_template_bytes(template, filename="even.docx")

    names = {item["name"] for item in placeholders}
    assert "default_hdr" in names
    assert "even_hdr" in names
    assert "body_var" in names


def test_extract_text_includes_first_page_header() -> None:
    """extract_docx_text must include text from first-page headers."""
    text = extract_docx_text(_build_first_page_header_only_bytes())
    assert "First page" in text


def test_iter_section_headers_skips_corrupted_header_gracefully() -> None:
    """If a header reference exists but the part is inaccessible, skip it."""
    doc = Document()
    doc.add_paragraph("Body")
    section = doc.sections[0]

    # Manually inject a header reference into the XML without a real part
    from lxml import etree
    ref = etree.SubElement(
        section._sectPr,
        qn("w:headerReference"),
    )
    ref.set(qn("w:type"), "default")
    ref.set(qn("r:id"), "rIdBogus")

    # This should not raise — the safety net catches the error
    headers = _iter_section_headers(section)
    # May return 0 or 1 depending on whether python-docx can resolve the ref
    assert isinstance(headers, list)


def test_iter_section_footers_returns_empty_for_no_references() -> None:
    """Sections without footer references must return an empty list."""
    doc = Document()
    doc.add_paragraph("Body")
    section = doc.sections[0]
    footers = _iter_section_footers(section)
    assert footers == []


def test_iter_section_headers_ignores_unknown_type() -> None:
    """Unknown w:type values in header references must be safely skipped."""
    doc = Document()
    doc.add_paragraph("Body")
    section = doc.sections[0]

    from lxml import etree
    ref = etree.SubElement(
        section._sectPr,
        qn("w:headerReference"),
    )
    ref.set(qn("w:type"), "unknown_type")
    ref.set(qn("r:id"), "rId999")

    headers = _iter_section_headers(section)
    assert headers == []
