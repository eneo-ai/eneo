from __future__ import annotations

import hashlib
import io
from datetime import datetime, timezone
from types import SimpleNamespace
from unittest.mock import AsyncMock, MagicMock
from uuid import uuid4

import pytest
from docx import Document

from eneo.authentication.principal_types import PrincipalType
from eneo.flows.domain.canonical_json_hash import canonical_json_bytes
from eneo.flows.domain.flow import (
    FlowRun,
    FlowRunStatus,
    FlowStepResult,
    FlowStepResultStatus,
)
from eneo.flows.domain.runtime import (
    RunExecutionState,
    RuntimeStep,
    StepExecutionOutput,
)
from eneo.flows.flow_api_error_code import FlowApiErrorCode
from eneo.flows.principal import FlowPrincipal
from eneo.flows.runtime import template_fill_runtime as template_fill_runtime_module
from eneo.flows.runtime.step_handlers.template_fill import TemplateFillStepHandler
from eneo.flows.runtime.template_fill_runtime import (
    TemplateFillRuntimeDeps,
    prepare_template_fill_step,
)
from eneo.flows.variable_resolver import FlowVariableResolver
from eneo.main.exceptions import (
    NotFoundException,
    TypedIOValidationException,
)


def _build_template_bytes() -> bytes:
    document = Document()
    document.add_paragraph("Titel: {{title}}")
    document.add_paragraph("Författare: {{author}}")
    document.add_paragraph("Sammanfattning: {{summary}}")
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _build_unicode_template_bytes() -> bytes:
    document = Document()
    document.add_paragraph("Ämne: {{ämne}}")
    document.add_paragraph("Sammanfattning: {{summary}}")
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _build_summary_only_template_bytes() -> bytes:
    document = Document()
    document.add_paragraph("Slutsats")
    document.add_paragraph("{{slutsats}}")
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _run() -> FlowRun:
    now = datetime.now(timezone.utc)
    return FlowRun(
        id=uuid4(),
        flow_id=uuid4(),
        flow_version=3,
        principal_type=PrincipalType.USER,
        principal_user_id=uuid4(),
        tenant_id=uuid4(),
        trace_id=uuid4(),
        status=FlowRunStatus.RUNNING,
        input_payload_json={
            "title": "Social medias påverkan",
            "author": "Anders Svensson",
        },
        created_at=now,
        updated_at=now,
    )


def _completed_result(*, run: FlowRun) -> FlowStepResult:
    now = datetime.now(timezone.utc)
    return FlowStepResult(
        id=uuid4(),
        flow_run_id=run.id,
        flow_id=run.flow_id,
        tenant_id=run.tenant_id,
        step_id=uuid4(),
        step_order=1,
        assistant_id=uuid4(),
        input_payload_json={"text": "ignored"},
        effective_prompt="prompt",
        output_payload_json={"text": "Detta är sammanfattningen."},
        model_parameters_json={},
        num_tokens_input=10,
        num_tokens_output=20,
        status=FlowStepResultStatus.COMPLETED,
        current_attempt_no=4,
        flow_step_execution_hash="hash",
        created_at=now,
        updated_at=now,
    )


def _state(*, result: FlowStepResult) -> RunExecutionState:
    return RunExecutionState(
        completed_by_order={1: result},
        prior_results=[result],
        assistant_cache={},
        json_mode_supported={},
        file_cache={},
        step_names_by_order={1: "Sammanfatta"},
    )


def _step(
    *,
    template_asset_id=None,
    template_checksum: str | None = "checksum",
) -> RuntimeStep:
    template_asset_id = template_asset_id or uuid4()
    return RuntimeStep(
        step_id=uuid4(),
        step_order=2,
        assistant_id=uuid4(),
        user_description="Assemble report",
        input_source="all_previous_steps",
        input_bindings=None,
        input_config=None,
        output_mode="template_fill",
        output_config={
            "template_asset_id": str(template_asset_id),
            "template_checksum": template_checksum,
            "bindings": {
                "title": "{{title}}",
                "author": "{{author}}",
                "summary": "{{step_1.output.text}}",
            },
        },
        output_type="docx",
    )


def _asset_step(
    *,
    template_asset_id,
    template_checksum: str | None = "checksum",
) -> RuntimeStep:
    return _step(
        template_asset_id=template_asset_id,
        template_checksum=template_checksum,
    )


def _logger() -> SimpleNamespace:
    return SimpleNamespace(
        info=lambda *args, **kwargs: None,
        debug=lambda *args, **kwargs: None,
        exception=lambda *args, **kwargs: None,
    )


async def execute_template_fill_step(
    *,
    step: RuntimeStep,
    run: FlowRun,
    state: RunExecutionState,
    deps: TemplateFillRuntimeDeps,
) -> StepExecutionOutput:
    handler = TemplateFillStepHandler(
        deps=deps,
        activate_resolved_input_edges=AsyncMock(),
    )
    result = await handler.execute(
        step=step,
        run=run,
        state=state,
        version_metadata=None,
        attempt_no=1,
    )
    return result.output


@pytest.mark.asyncio
async def test_prepare_template_fill_step_static_binding_adds_no_resolved_input_edge() -> (
    None
):
    run = _run()
    result = _completed_result(run=run)
    state = _state(result=result)
    template_file_id = uuid4()
    file_repo = AsyncMock()
    file_repo.get_list_by_id_and_tenant.return_value = [
        SimpleNamespace(
            id=template_file_id,
            name="template.docx",
            checksum="checksum",
            blob=_build_template_bytes(),
        )
    ]
    template_asset_repo = AsyncMock()
    template_asset_repo.get.return_value = SimpleNamespace(file_id=template_file_id)
    deps = TemplateFillRuntimeDeps(
        variable_resolver=FlowVariableResolver(),
        file_repo=file_repo,
        template_asset_repo=template_asset_repo,
        principal=FlowPrincipal.from_run(run),
        logger=_logger(),
    )
    step = _step()
    step.output_config["bindings"] = {"title": "Statisk titel"}
    step.output_config["placeholders"] = ["title"]

    prepared = await prepare_template_fill_step(
        step=step,
        run=run,
        state=state,
        deps=deps,
    )

    assert prepared.resolved_bindings == {"title": "Statisk titel"}
    assert prepared.resolved_input_edges == ()


@pytest.mark.asyncio
async def test_prepare_template_fill_step_records_exact_dynamic_binding_source() -> (
    None
):
    run = _run()
    result = _completed_result(run=run)
    state = _state(result=result)
    template_file_id = uuid4()
    file_repo = AsyncMock()
    file_repo.get_list_by_id_and_tenant.return_value = [
        SimpleNamespace(
            id=template_file_id,
            name="template.docx",
            checksum="checksum",
            blob=_build_template_bytes(),
        )
    ]
    template_asset_repo = AsyncMock()
    template_asset_repo.get.return_value = SimpleNamespace(file_id=template_file_id)
    deps = TemplateFillRuntimeDeps(
        variable_resolver=FlowVariableResolver(),
        file_repo=file_repo,
        template_asset_repo=template_asset_repo,
        principal=FlowPrincipal.from_run(run),
        logger=_logger(),
    )
    step = _step()
    step.output_config["bindings"] = {
        "summary": "{{step_1.output.text}}",
    }
    step.output_config["placeholders"] = ["summary"]

    prepared = await prepare_template_fill_step(
        step=step,
        run=run,
        state=state,
        deps=deps,
    )

    assert prepared.resolved_bindings == {"summary": "Detta är sammanfattningen."}
    assert len(prepared.resolved_input_edges) == 1
    edge = prepared.resolved_input_edges[0]
    assert edge.binding_ref == "output_config.bindings.summary:step_1.output.text"
    assert edge.source.kind == "step_result"
    assert edge.source.source_step_id == result.step_id
    assert edge.source.source_attempt_no == 4
    assert edge.source.selector.path == ("output", "text")
    selected = "Detta är sammanfattningen.".encode("utf-8")
    assert edge.selection.encoding == "utf8"
    assert edge.selection.byte_size == len(selected)
    assert edge.selection.sha256 == hashlib.sha256(selected).hexdigest()


@pytest.mark.asyncio
async def test_template_fill_activation_failure_prevents_render_and_persist(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    run = _run()
    result = _completed_result(run=run)
    state = _state(result=result)
    template_file_id = uuid4()
    file_repo = AsyncMock()
    file_repo.get_list_by_id_and_tenant.return_value = [
        SimpleNamespace(
            id=template_file_id,
            name="template.docx",
            checksum="checksum",
            blob=_build_template_bytes(),
        )
    ]
    template_asset_repo = AsyncMock()
    template_asset_repo.get.return_value = SimpleNamespace(file_id=template_file_id)
    deps = TemplateFillRuntimeDeps(
        variable_resolver=FlowVariableResolver(),
        file_repo=file_repo,
        template_asset_repo=template_asset_repo,
        principal=FlowPrincipal.from_run(run),
        logger=_logger(),
    )
    render_docx_template = MagicMock()
    monkeypatch.setattr(
        template_fill_runtime_module,
        "render_docx_template",
        render_docx_template,
    )
    handler = TemplateFillStepHandler(
        deps=deps,
        activate_resolved_input_edges=AsyncMock(
            side_effect=RuntimeError("activation failed")
        ),
    )

    with pytest.raises(RuntimeError, match="activation failed"):
        await handler.execute(
            step=_step(),
            run=run,
            state=state,
            version_metadata=None,
            attempt_no=1,
        )

    render_docx_template.assert_not_called()
    file_repo.add.assert_not_awaited()


@pytest.mark.asyncio
async def test_execute_template_fill_step_renders_and_persists_docx() -> None:
    run = _run()
    result = _completed_result(run=run)
    state = _state(result=result)
    template_asset_id = uuid4()
    template_file_id = uuid4()
    step = _step(template_asset_id=template_asset_id)
    file_repo = AsyncMock()
    file_repo.get_list_by_id_and_tenant.return_value = [
        SimpleNamespace(
            id=template_file_id,
            name="template.docx",
            checksum="checksum",
            blob=_build_template_bytes(),
        )
    ]
    file_repo.add.return_value = SimpleNamespace(id=uuid4())
    template_asset_repo = AsyncMock()
    template_asset_repo.get.return_value = SimpleNamespace(file_id=template_file_id)
    deps = TemplateFillRuntimeDeps(
        variable_resolver=FlowVariableResolver(),
        file_repo=file_repo,
        template_asset_repo=template_asset_repo,
        principal=FlowPrincipal.from_run(run),
        logger=_logger(),
    )

    output = await execute_template_fill_step(
        step=step,
        run=run,
        state=state,
        deps=deps,
    )

    assert output.persisted_text == (
        "## title\n\nSocial medias påverkan\n\n"
        "## author\n\nAnders Svensson\n\n"
        "## summary\n\nDetta är sammanfattningen."
    )
    assert output.artifacts == [
        {
            "file_id": str(file_repo.add.return_value.id),
            "name": "step_2_output.docx",
            "mimetype": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "size": len(file_repo.add.await_args.args[0].blob),
            "checksum": file_repo.add.await_args.args[0].checksum,
            "file_type": "document",
        }
    ]
    assert output.model_parameters_json["mode"] == "template_fill"
    assert "Detta är sammanfattningen." in output.full_text
    assert output.output_payload_extensions == {
        "template_fill_debug": {
            "rendered_docx_text_raw": output.full_text,
            "summary_mode": "resolved_bindings",
            "placeholder_count": 3,
        },
        "template_provenance": {
            "template_name": "template.docx",
            "template_asset_id": str(template_asset_id),
            "template_file_id": str(template_file_id),
            "template_checksum": "checksum",
            "published_flow_version": run.flow_version,
        },
    }
    assert output.num_tokens_input == 0
    assert output.num_tokens_output == 0
    assert '"title": "Social medias påverkan"' in output.input_text
    file_repo.add.assert_awaited_once()
    template_asset_repo.get.assert_awaited_once_with(
        asset_id=template_asset_id,
        tenant_id=run.tenant_id,
    )


@pytest.mark.asyncio
async def test_execute_template_fill_step_loads_template_asset_by_tenant() -> None:
    run = _run()
    result = _completed_result(run=run)
    state = _state(result=result)
    template_asset_id = uuid4()
    template_file_id = uuid4()
    file_repo = AsyncMock()
    file_repo.get_list_by_id_and_tenant.return_value = [
        SimpleNamespace(
            id=template_file_id,
            name="template.docx",
            checksum="checksum",
            blob=_build_template_bytes(),
        )
    ]
    file_repo.add.return_value = SimpleNamespace(id=uuid4())
    template_asset_repo = AsyncMock()
    template_asset_repo.get.return_value = SimpleNamespace(file_id=template_file_id)
    deps = TemplateFillRuntimeDeps(
        variable_resolver=FlowVariableResolver(),
        file_repo=file_repo,
        template_asset_repo=template_asset_repo,
        principal=FlowPrincipal.from_run(run),
        logger=_logger(),
    )

    output = await execute_template_fill_step(
        step=_asset_step(
            template_asset_id=template_asset_id,
        ),
        run=run,
        state=state,
        deps=deps,
    )

    template_asset_repo.get.assert_awaited_once_with(
        asset_id=template_asset_id,
        tenant_id=run.tenant_id,
    )
    file_repo.get_list_by_id_and_tenant.assert_awaited_once_with(
        ids=[template_file_id],
        tenant_id=run.tenant_id,
        include_transcription=False,
    )
    assert output.output_payload_extensions["template_provenance"] == {
        "template_name": "template.docx",
        "template_asset_id": str(template_asset_id),
        "template_file_id": str(template_file_id),
        "template_checksum": "checksum",
        "published_flow_version": run.flow_version,
    }


@pytest.mark.asyncio
async def test_execute_template_fill_step_resolves_asset_id_only_config() -> None:
    run = _run()
    result = _completed_result(run=run)
    state = _state(result=result)
    template_asset_id = uuid4()
    asset_file_id = uuid4()
    step = _asset_step(
        template_asset_id=template_asset_id,
    )
    file_repo = AsyncMock()
    file_repo.get_list_by_id_and_tenant.return_value = [
        SimpleNamespace(
            id=asset_file_id,
            name="template.docx",
            checksum="checksum",
            blob=_build_template_bytes(),
        )
    ]
    file_repo.add.return_value = SimpleNamespace(id=uuid4())
    template_asset_repo = AsyncMock()
    template_asset_repo.get.return_value = SimpleNamespace(file_id=asset_file_id)
    deps = TemplateFillRuntimeDeps(
        variable_resolver=FlowVariableResolver(),
        file_repo=file_repo,
        template_asset_repo=template_asset_repo,
        principal=FlowPrincipal.from_run(run),
        logger=_logger(),
    )

    output = await execute_template_fill_step(
        step=step,
        run=run,
        state=state,
        deps=deps,
    )

    template_asset_repo.get.assert_awaited_once_with(
        asset_id=template_asset_id,
        tenant_id=run.tenant_id,
    )
    file_repo.get_list_by_id_and_tenant.assert_awaited_once_with(
        ids=[asset_file_id],
        tenant_id=run.tenant_id,
        include_transcription=False,
    )
    assert output.output_payload_extensions["template_provenance"] == {
        "template_name": "template.docx",
        "template_asset_id": str(template_asset_id),
        "template_file_id": str(asset_file_id),
        "template_checksum": "checksum",
        "published_flow_version": run.flow_version,
    }
    assert output.model_parameters_json["template_asset_id"] == str(template_asset_id)
    assert output.model_parameters_json["template_file_id"] == str(asset_file_id)


@pytest.mark.asyncio
async def test_execute_template_fill_step_reports_missing_template_asset() -> None:
    run = _run()
    result = _completed_result(run=run)
    state = _state(result=result)
    template_asset_id = uuid4()
    template_asset_repo = AsyncMock()
    template_asset_repo.get.side_effect = NotFoundException("missing asset")
    file_repo = AsyncMock()
    deps = TemplateFillRuntimeDeps(
        variable_resolver=FlowVariableResolver(),
        file_repo=file_repo,
        template_asset_repo=template_asset_repo,
        principal=FlowPrincipal.from_run(run),
        logger=_logger(),
    )

    with pytest.raises(TypedIOValidationException) as exc_info:
        await execute_template_fill_step(
            step=_asset_step(
                template_asset_id=template_asset_id,
            ),
            run=run,
            state=state,
            deps=deps,
        )

    assert exc_info.value.code == FlowApiErrorCode.TEMPLATE_NOT_ACCESSIBLE.value
    file_repo.get_list_by_id_and_tenant.assert_not_awaited()


@pytest.mark.asyncio
async def test_execute_template_fill_step_reports_missing_template_asset_file() -> None:
    run = _run()
    result = _completed_result(run=run)
    state = _state(result=result)
    template_asset_id = uuid4()
    template_file_id = uuid4()
    template_asset_repo = AsyncMock()
    template_asset_repo.get.return_value = SimpleNamespace(file_id=template_file_id)
    file_repo = AsyncMock()
    file_repo.get_list_by_id_and_tenant.return_value = []
    deps = TemplateFillRuntimeDeps(
        variable_resolver=FlowVariableResolver(),
        file_repo=file_repo,
        template_asset_repo=template_asset_repo,
        principal=FlowPrincipal.from_run(run),
        logger=_logger(),
    )

    with pytest.raises(TypedIOValidationException) as exc_info:
        await execute_template_fill_step(
            step=_asset_step(
                template_asset_id=template_asset_id,
            ),
            run=run,
            state=state,
            deps=deps,
        )

    assert exc_info.value.code == FlowApiErrorCode.TEMPLATE_NOT_ACCESSIBLE.value


@pytest.mark.asyncio
async def test_execute_template_fill_step_rejects_template_file_id_identity() -> None:
    run = _run()
    result = _completed_result(run=run)
    state = _state(result=result)
    template_asset_id = uuid4()
    template_file_id = uuid4()
    step = _asset_step(
        template_asset_id=template_asset_id,
    )
    step.output_config["template_file_id"] = str(template_file_id)
    template_asset_repo = AsyncMock()
    file_repo = AsyncMock()
    deps = TemplateFillRuntimeDeps(
        variable_resolver=FlowVariableResolver(),
        file_repo=file_repo,
        template_asset_repo=template_asset_repo,
        principal=FlowPrincipal.from_run(run),
        logger=_logger(),
    )

    with pytest.raises(
        TypedIOValidationException, match="template_file_id is not supported"
    ):
        await execute_template_fill_step(
            step=step,
            run=run,
            state=state,
            deps=deps,
        )

    template_asset_repo.get.assert_not_awaited()
    file_repo.get_list_by_id_and_tenant.assert_not_awaited()


@pytest.mark.asyncio
async def test_execute_template_fill_step_reports_missing_template_asset_content() -> (
    None
):
    run = _run()
    result = _completed_result(run=run)
    state = _state(result=result)
    template_asset_id = uuid4()
    template_file_id = uuid4()
    template_asset_repo = AsyncMock()
    template_asset_repo.get.return_value = SimpleNamespace(file_id=template_file_id)
    file_repo = AsyncMock()
    file_repo.get_list_by_id_and_tenant.return_value = [
        SimpleNamespace(
            id=template_file_id,
            name="template.docx",
            checksum="checksum",
            blob=None,
        )
    ]
    deps = TemplateFillRuntimeDeps(
        variable_resolver=FlowVariableResolver(),
        file_repo=file_repo,
        template_asset_repo=template_asset_repo,
        principal=FlowPrincipal.from_run(run),
        logger=_logger(),
    )

    with pytest.raises(TypedIOValidationException) as exc_info:
        await execute_template_fill_step(
            step=_asset_step(
                template_asset_id=template_asset_id,
            ),
            run=run,
            state=state,
            deps=deps,
        )

    assert exc_info.value.code == FlowApiErrorCode.TEMPLATE_MISSING_CONTENT.value


@pytest.mark.asyncio
async def test_execute_template_fill_step_reports_template_asset_checksum_drift() -> (
    None
):
    run = _run()
    result = _completed_result(run=run)
    state = _state(result=result)
    template_asset_id = uuid4()
    template_file_id = uuid4()
    template_asset_repo = AsyncMock()
    template_asset_repo.get.return_value = SimpleNamespace(file_id=template_file_id)
    file_repo = AsyncMock()
    file_repo.get_list_by_id_and_tenant.return_value = [
        SimpleNamespace(
            id=template_file_id,
            name="template.docx",
            checksum="other-checksum",
            blob=_build_template_bytes(),
        )
    ]
    deps = TemplateFillRuntimeDeps(
        variable_resolver=FlowVariableResolver(),
        file_repo=file_repo,
        template_asset_repo=template_asset_repo,
        principal=FlowPrincipal.from_run(run),
        logger=_logger(),
    )

    with pytest.raises(TypedIOValidationException) as exc_info:
        await execute_template_fill_step(
            step=_asset_step(
                template_asset_id=template_asset_id,
                template_checksum="expected-checksum",
            ),
            run=run,
            state=state,
            deps=deps,
        )

    assert exc_info.value.code == "typed_io_template_checksum_mismatch"


@pytest.mark.asyncio
async def test_execute_template_fill_step_rejects_checksum_drift() -> None:
    run = _run()
    result = _completed_result(run=run)
    state = _state(result=result)
    template_file_id = uuid4()
    file_repo = AsyncMock()
    file_repo.get_list_by_id_and_tenant.return_value = [
        SimpleNamespace(
            id=template_file_id,
            name="template.docx",
            checksum="other-checksum",
            blob=_build_template_bytes(),
        )
    ]
    deps = TemplateFillRuntimeDeps(
        variable_resolver=FlowVariableResolver(),
        file_repo=file_repo,
        template_asset_repo=AsyncMock(),
        principal=FlowPrincipal.from_run(run),
        logger=_logger(),
    )

    with pytest.raises(TypedIOValidationException, match="checksum"):
        await execute_template_fill_step(
            step=_step(template_checksum="expected-checksum"),
            run=run,
            state=state,
            deps=deps,
        )

    file_repo.add.assert_not_called()


@pytest.mark.asyncio
async def test_execute_template_fill_step_allows_explicit_empty_binding() -> None:
    run = _run()
    result = _completed_result(run=run)
    state = _state(result=result)
    template_file_id = uuid4()
    file_repo = AsyncMock()
    file_repo.get_list_by_id_and_tenant.return_value = [
        SimpleNamespace(
            id=template_file_id,
            name="template.docx",
            checksum="checksum",
            blob=_build_template_bytes(),
        )
    ]
    file_repo.add.return_value = SimpleNamespace(id=uuid4())
    deps = TemplateFillRuntimeDeps(
        variable_resolver=FlowVariableResolver(),
        file_repo=file_repo,
        template_asset_repo=AsyncMock(),
        principal=FlowPrincipal.from_run(run),
        logger=_logger(),
    )

    step = _step()
    step.output_config["bindings"]["author"] = ""

    output = await execute_template_fill_step(
        step=step,
        run=run,
        state=state,
        deps=deps,
    )

    assert '"author": ""' in output.input_text
    assert "\n\n## author\n\n## summary" in output.persisted_text


@pytest.mark.asyncio
async def test_execute_template_fill_step_strips_duplicate_leading_heading_from_summary_only() -> (
    None
):
    run = _run()
    now = datetime.now(timezone.utc)
    result = FlowStepResult(
        id=uuid4(),
        flow_run_id=run.id,
        flow_id=run.flow_id,
        tenant_id=run.tenant_id,
        step_id=uuid4(),
        step_order=1,
        assistant_id=uuid4(),
        input_payload_json={"text": "ignored"},
        effective_prompt="prompt",
        output_payload_json={"text": "Slutsats:\n\nDet här är själva brödtexten."},
        model_parameters_json={},
        num_tokens_input=10,
        num_tokens_output=20,
        status=FlowStepResultStatus.COMPLETED,
        current_attempt_no=4,
        flow_step_execution_hash="hash",
        created_at=now,
        updated_at=now,
    )
    state = RunExecutionState(
        completed_by_order={1: result},
        prior_results=[result],
        assistant_cache={},
        json_mode_supported={},
        file_cache={},
        step_names_by_order={1: "slutsats"},
    )
    template_file_id = uuid4()
    file_repo = AsyncMock()
    file_repo.get_list_by_id_and_tenant.return_value = [
        SimpleNamespace(
            id=template_file_id,
            name="template.docx",
            checksum="checksum",
            blob=_build_summary_only_template_bytes(),
        )
    ]
    file_repo.add.return_value = SimpleNamespace(id=uuid4())
    deps = TemplateFillRuntimeDeps(
        variable_resolver=FlowVariableResolver(),
        file_repo=file_repo,
        template_asset_repo=AsyncMock(),
        principal=FlowPrincipal.from_run(run),
        logger=_logger(),
    )
    step = _step()
    step.output_config["bindings"] = {"slutsats": "{{step_1.output.text}}"}
    step.output_config["placeholders"] = ["slutsats"]

    output = await execute_template_fill_step(
        step=step,
        run=run,
        state=state,
        deps=deps,
    )

    assert output.persisted_text == "## slutsats\n\nDet här är själva brödtexten."
    assert "Slutsats:\n\nDet här är själva brödtexten." in output.full_text


@pytest.mark.asyncio
async def test_execute_template_fill_step_reports_failed_upstream_step_clearly() -> (
    None
):
    run = _run()
    now = datetime.now(timezone.utc)
    failed_result = FlowStepResult(
        id=uuid4(),
        flow_run_id=run.id,
        flow_id=run.flow_id,
        tenant_id=run.tenant_id,
        step_id=uuid4(),
        step_order=1,
        assistant_id=uuid4(),
        input_payload_json={"text": "ignored"},
        effective_prompt="prompt",
        output_payload_json={},
        model_parameters_json={},
        num_tokens_input=1,
        num_tokens_output=1,
        status=FlowStepResultStatus.FAILED,
        error_message="boom",
        flow_step_execution_hash="hash",
        created_at=now,
        updated_at=now,
    )
    state = RunExecutionState(
        completed_by_order={},
        prior_results=[failed_result],
        assistant_cache={},
        json_mode_supported={},
        file_cache={},
        step_names_by_order={1: "Sammanfatta"},
    )
    template_file_id = uuid4()
    file_repo = AsyncMock()
    file_repo.get_list_by_id_and_tenant.return_value = [
        SimpleNamespace(
            id=template_file_id,
            name="template.docx",
            checksum="checksum",
            blob=_build_template_bytes(),
        )
    ]
    deps = TemplateFillRuntimeDeps(
        variable_resolver=FlowVariableResolver(),
        file_repo=file_repo,
        template_asset_repo=AsyncMock(),
        principal=FlowPrincipal.from_run(run),
        logger=_logger(),
    )

    with pytest.raises(TypedIOValidationException, match="failed earlier in the run"):
        await execute_template_fill_step(
            step=_step(),
            run=run,
            state=state,
            deps=deps,
        )


@pytest.mark.asyncio
async def test_execute_template_fill_step_reports_missing_template_blob_clearly() -> (
    None
):
    run = _run()
    result = _completed_result(run=run)
    state = _state(result=result)
    template_file_id = uuid4()
    file_repo = AsyncMock()
    file_repo.get_list_by_id_and_tenant.return_value = [
        SimpleNamespace(
            id=template_file_id,
            name="template.docx",
            checksum="checksum",
            blob=None,
        )
    ]
    deps = TemplateFillRuntimeDeps(
        variable_resolver=FlowVariableResolver(),
        file_repo=file_repo,
        template_asset_repo=AsyncMock(),
        principal=FlowPrincipal.from_run(run),
        logger=_logger(),
    )

    with pytest.raises(
        TypedIOValidationException,
        match="could not be read because the saved file content is missing",
    ):
        await execute_template_fill_step(
            step=_step(),
            run=run,
            state=state,
            deps=deps,
        )


@pytest.mark.asyncio
async def test_execute_template_fill_step_reports_generated_document_save_failure() -> (
    None
):
    run = _run()
    result = _completed_result(run=run)
    state = _state(result=result)
    template_file_id = uuid4()
    file_repo = AsyncMock()
    file_repo.get_list_by_id_and_tenant.return_value = [
        SimpleNamespace(
            id=template_file_id,
            name="template.docx",
            checksum="checksum",
            blob=_build_template_bytes(),
        )
    ]
    file_repo.add.side_effect = RuntimeError("storage unavailable")
    deps = TemplateFillRuntimeDeps(
        variable_resolver=FlowVariableResolver(),
        file_repo=file_repo,
        template_asset_repo=AsyncMock(),
        principal=FlowPrincipal.from_run(run),
        logger=_logger(),
    )

    with pytest.raises(TypedIOValidationException, match="saving the generated DOCX"):
        await execute_template_fill_step(
            step=_step(),
            run=run,
            state=state,
            deps=deps,
        )


@pytest.mark.asyncio
async def test_execute_template_fill_step_reports_generated_document_read_failure(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    run = _run()
    result = _completed_result(run=run)
    state = _state(result=result)
    template_file_id = uuid4()
    file_repo = AsyncMock()
    file_repo.get_list_by_id_and_tenant.return_value = [
        SimpleNamespace(
            id=template_file_id,
            name="template.docx",
            checksum="checksum",
            blob=_build_template_bytes(),
        )
    ]
    file_repo.add.return_value = SimpleNamespace(id=uuid4())
    monkeypatch.setattr(
        template_fill_runtime_module,
        "extract_docx_text",
        lambda blob: (_ for _ in ()).throw(RuntimeError("document parse failed")),
    )
    deps = TemplateFillRuntimeDeps(
        variable_resolver=FlowVariableResolver(),
        file_repo=file_repo,
        template_asset_repo=AsyncMock(),
        principal=FlowPrincipal.from_run(run),
        logger=_logger(),
    )

    with pytest.raises(TypedIOValidationException, match="reading the generated DOCX"):
        await execute_template_fill_step(
            step=_step(),
            run=run,
            state=state,
            deps=deps,
        )


@pytest.mark.asyncio
async def test_execute_template_fill_step_rejects_template_file_id_key() -> None:
    run = _run()
    result = _completed_result(run=run)
    state = _state(result=result)
    step = _step()
    step.output_config["template_file_id"] = str(uuid4())
    deps = TemplateFillRuntimeDeps(
        variable_resolver=FlowVariableResolver(),
        file_repo=AsyncMock(),
        template_asset_repo=AsyncMock(),
        principal=FlowPrincipal.from_run(run),
        logger=_logger(),
    )

    with pytest.raises(
        TypedIOValidationException, match="template_file_id is not supported"
    ):
        await execute_template_fill_step(
            step=step,
            run=run,
            state=state,
            deps=deps,
        )


@pytest.mark.asyncio
async def test_execute_template_fill_step_rejects_non_string_binding_values() -> None:
    run = _run()
    result = _completed_result(run=run)
    state = _state(result=result)
    step = _step()
    step.output_config["bindings"]["summary"] = {"bad": "value"}
    deps = TemplateFillRuntimeDeps(
        variable_resolver=FlowVariableResolver(),
        file_repo=AsyncMock(),
        template_asset_repo=AsyncMock(),
        principal=FlowPrincipal.from_run(run),
        logger=_logger(),
    )

    with pytest.raises(TypedIOValidationException, match="must be a string expression"):
        await execute_template_fill_step(
            step=step,
            run=run,
            state=state,
            deps=deps,
        )


@pytest.mark.asyncio
async def test_execute_template_fill_step_reports_missing_published_template_file() -> (
    None
):
    run = _run()
    result = _completed_result(run=run)
    state = _state(result=result)
    file_repo = AsyncMock()
    file_repo.get_list_by_id_and_tenant.return_value = []
    deps = TemplateFillRuntimeDeps(
        variable_resolver=FlowVariableResolver(),
        file_repo=file_repo,
        template_asset_repo=AsyncMock(),
        principal=FlowPrincipal.from_run(run),
        logger=_logger(),
    )

    with pytest.raises(
        TypedIOValidationException, match="asset file is no longer available"
    ):
        await execute_template_fill_step(
            step=_step(),
            run=run,
            state=state,
            deps=deps,
        )


@pytest.mark.asyncio
async def test_execute_template_fill_step_reports_render_stage_failure(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    run = _run()
    result = _completed_result(run=run)
    state = _state(result=result)
    template_file_id = uuid4()
    file_repo = AsyncMock()
    file_repo.get_list_by_id_and_tenant.return_value = [
        SimpleNamespace(
            id=template_file_id,
            name="template.docx",
            checksum="checksum",
            blob=_build_template_bytes(),
        )
    ]
    monkeypatch.setattr(
        template_fill_runtime_module,
        "render_docx_template",
        lambda **kwargs: (_ for _ in ()).throw(RuntimeError("render broke")),
    )
    deps = TemplateFillRuntimeDeps(
        variable_resolver=FlowVariableResolver(),
        file_repo=file_repo,
        template_asset_repo=AsyncMock(),
        principal=FlowPrincipal.from_run(run),
        logger=_logger(),
    )

    with pytest.raises(TypedIOValidationException, match="rendering the DOCX template"):
        await execute_template_fill_step(
            step=_step(),
            run=run,
            state=state,
            deps=deps,
        )


@pytest.mark.asyncio
async def test_execute_template_fill_step_supports_datum_system_variable() -> None:
    run = _run()
    result = _completed_result(run=run)
    state = _state(result=result)
    template_file_id = uuid4()
    file_repo = AsyncMock()
    file_repo.get_list_by_id_and_tenant.return_value = [
        SimpleNamespace(
            id=template_file_id,
            name="template.docx",
            checksum="checksum",
            blob=_build_template_bytes(),
        )
    ]
    file_repo.add.return_value = SimpleNamespace(id=uuid4())
    deps = TemplateFillRuntimeDeps(
        variable_resolver=FlowVariableResolver(),
        file_repo=file_repo,
        template_asset_repo=AsyncMock(),
        principal=FlowPrincipal.from_run(run),
        logger=_logger(),
    )
    step = _step()
    step.output_config["bindings"]["author"] = "{{datum}}"

    output = await execute_template_fill_step(
        step=step,
        run=run,
        state=state,
        deps=deps,
    )

    assert '"author": "' in output.input_text
    assert output.input_text.count("-") >= 2
    assert "## author\n\n" in output.persisted_text


@pytest.mark.asyncio
async def test_execute_template_fill_step_formats_json_binding_in_summary() -> None:
    run = _run()
    result = _completed_result(run=run)
    result = result.model_copy(
        update={
            "output_payload_json": {
                "text": '{"status":"preview"}',
                "structured": {"status": "approved", "score": 5},
            }
        },
        deep=True,
    )
    state = _state(result=result)
    template_file_id = uuid4()
    file_repo = AsyncMock()
    file_repo.get_list_by_id_and_tenant.return_value = [
        SimpleNamespace(
            id=template_file_id,
            name="template.docx",
            checksum="checksum",
            blob=_build_template_bytes(),
        )
    ]
    file_repo.add.return_value = SimpleNamespace(id=uuid4())
    deps = TemplateFillRuntimeDeps(
        variable_resolver=FlowVariableResolver(),
        file_repo=file_repo,
        template_asset_repo=AsyncMock(),
        principal=FlowPrincipal.from_run(run),
        logger=_logger(),
    )
    step = _step()
    step.output_config["bindings"]["summary"] = "{{step_1.output.structured}}"

    prepared = await prepare_template_fill_step(
        step=step,
        run=run,
        state=state,
        deps=deps,
    )
    structured_edge = next(
        edge
        for edge in prepared.resolved_input_edges
        if edge.binding_ref.endswith(":step_1.output.structured")
    )
    selected = canonical_json_bytes({"status": "approved", "score": 5})
    assert structured_edge.selection.encoding == "canonical_json"
    assert structured_edge.selection.byte_size == len(selected)
    assert structured_edge.selection.sha256 == hashlib.sha256(selected).hexdigest()

    output = await execute_template_fill_step(
        step=step, run=run, state=state, deps=deps
    )

    assert '## summary\n\n{"status": "approved", "score": 5}' in output.persisted_text


@pytest.mark.asyncio
async def test_execute_template_fill_step_supports_unicode_placeholder_names() -> None:
    run = _run()
    result = _completed_result(run=run)
    state = _state(result=result)
    template_file_id = uuid4()
    file_repo = AsyncMock()
    file_repo.get_list_by_id_and_tenant.return_value = [
        SimpleNamespace(
            id=template_file_id,
            name="template.docx",
            checksum="checksum",
            blob=_build_unicode_template_bytes(),
        )
    ]
    file_repo.add.return_value = SimpleNamespace(id=uuid4())
    deps = TemplateFillRuntimeDeps(
        variable_resolver=FlowVariableResolver(),
        file_repo=file_repo,
        template_asset_repo=AsyncMock(),
        principal=FlowPrincipal.from_run(run),
        logger=_logger(),
    )
    step = _step()
    step.output_config["placeholders"] = ["ämne", "summary"]
    step.output_config["bindings"] = {
        "ämne": "{{title}}",
        "summary": "{{step_1.output.text}}",
    }

    output = await execute_template_fill_step(
        step=step, run=run, state=state, deps=deps
    )

    assert output.persisted_text.startswith(
        "## ämne\n\nSocial medias påverkan\n\n## summary"
    )


@pytest.mark.asyncio
async def test_execute_template_fill_step_preserves_long_summary_in_persisted_text() -> (
    None
):
    run = _run()
    result = _completed_result(run=run)
    result = result.model_copy(
        update={"output_payload_json": {"text": "L" * 10000}},
        deep=True,
    )
    state = _state(result=result)
    template_file_id = uuid4()
    file_repo = AsyncMock()
    file_repo.get_list_by_id_and_tenant.return_value = [
        SimpleNamespace(
            id=template_file_id,
            name="template.docx",
            checksum="checksum",
            blob=_build_template_bytes(),
        )
    ]
    file_repo.add.return_value = SimpleNamespace(id=uuid4())
    deps = TemplateFillRuntimeDeps(
        variable_resolver=FlowVariableResolver(),
        file_repo=file_repo,
        template_asset_repo=AsyncMock(),
        principal=FlowPrincipal.from_run(run),
        logger=_logger(),
    )

    output = await execute_template_fill_step(
        step=_step(),
        run=run,
        state=state,
        deps=deps,
    )

    assert "## summary" in output.persisted_text
    assert "L" * 10000 in output.persisted_text
