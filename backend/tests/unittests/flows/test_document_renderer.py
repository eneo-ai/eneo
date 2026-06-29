"""Tests for DocumentRenderService PDF/DOCX generation."""

from __future__ import annotations

import io
from collections.abc import Sequence
from unittest.mock import patch

import pdfplumber
import pytest
from docx import Document

from intric.flows.runtime.document_rendering.blocks import DocumentBlock, InlineTextRun
from intric.flows.runtime.document_rendering.html_blocks import (
    blocks_to_html_document,
)
from intric.flows.runtime.document_rendering.limits import DocumentRenderLimits
from intric.flows.runtime.document_rendering.markdown_blocks import (
    parse_markdown_blocks,
)
from intric.flows.runtime.document_rendering.renderers import RenderedDocument
from intric.flows.runtime.document_rendering.service import (
    DocumentRenderService,
    default_document_render_service,
)
from intric.flows.runtime.document_rendering.weasyprint_renderer import (
    WeasyPrintDocumentRenderer,
    _deny_external_fetch,
    configure_weasyprint_dependency_logging,
)
from intric.main.exceptions import TypedIOValidationException

_render_service = default_document_render_service()

# --- PDF rendering ---


def test_render_pdf_valid_blob():
    blob, mimetype, filename = _render_service.render_document("Hello world", "pdf", step_order=1)
    assert isinstance(blob, bytes)
    assert len(blob) > 0
    assert blob[:5] == b"%PDF-"


def test_render_pdf_correct_mime():
    _, mimetype, _ = _render_service.render_document("Test", "pdf", step_order=1)
    assert mimetype == "application/pdf"


def test_render_pdf_filename_pattern():
    _, _, filename = _render_service.render_document("Test", "pdf", step_order=3)
    assert filename == "step_3_output.pdf"


def test_render_pdf_markdown_headings_and_lists_as_readable_document():
    text = "# Titel\n\n## Sammanfattning\n\n- punkt ett\n- punkt två\n\n1. nästa steg"

    blob, _, _ = _render_service.render_document(text, "pdf", step_order=1)

    with pdfplumber.open(io.BytesIO(blob)) as pdf:
        page_text = "\n".join(page.extract_text() or "" for page in pdf.pages)
    assert "Titel" in page_text
    assert "Sammanfattning" in page_text
    assert "punkt ett" in page_text
    assert "# Titel" not in page_text
    assert "## Sammanfattning" not in page_text


def test_render_pdf_markdown_table_outputs_cells_without_separator_row():
    text = "| Namn | Värde |\n| --- | --- |\n| Kommun | Sundsvall |"

    blob, _, _ = _render_service.render_document(text, "pdf", step_order=1)

    with pdfplumber.open(io.BytesIO(blob)) as pdf:
        page_text = "\n".join(page.extract_text() or "" for page in pdf.pages)
    assert "Namn" in page_text
    assert "Sundsvall" in page_text
    assert "---" not in page_text


def test_markdown_blocks_normalize_common_markdown_control_syntax():
    blocks = parse_markdown_blocks(
        [
            "Medarbetarsamtal – Sammanfattning & Plan",
            "",
            "**Medarbetare:** [Leona](https://example.com)",
            "---",
            "1) **Sammanfattning**",
            "- Träna `skrivande` hemma",
        ]
    )

    assert [(block.kind, block.text, block.items) for block in blocks] == [
        ("paragraph", "Medarbetarsamtal – Sammanfattning & Plan", ()),
        ("paragraph", "Medarbetare: Leona", ()),
        ("empty", "", ()),
        ("numbered_list", "", ("Sammanfattning",)),
        ("bullet_list", "", ("Träna skrivande hemma",)),
    ]


def test_render_pdf_markdown_inline_syntax_as_readable_text():
    text = (
        "Medarbetarsamtal – Sammanfattning & Plan\n\n"
        "**Medarbetare:** Leona\n"
        "**Datum:** [Fyll i datum]\n\n"
        "---\n\n"
        "1) **Sammanfattning**\n"
        "Leona visar tydliga framsteg inom **svenska** och `matematik`."
    )

    blob, _, _ = _render_service.render_document(text, "pdf", step_order=1)

    with pdfplumber.open(io.BytesIO(blob)) as pdf:
        page_text = "\n".join(page.extract_text() or "" for page in pdf.pages)
    assert "Medarbetare: Leona" in page_text
    assert "Datum: [Fyll i datum]" in page_text
    assert "1. Sammanfattning" in page_text
    assert "svenska" in page_text
    assert "matematik" in page_text
    assert "**" not in page_text
    assert "---" not in page_text
    assert "`" not in page_text


def test_markdown_blocks_keep_code_block_content_verbatim():
    blocks = parse_markdown_blocks(
        [
            "```",
            "**not bold**",
            "---",
            "```",
        ]
    )

    assert blocks == [DocumentBlock(kind="code", text="**not bold**\n---")]


def test_markdown_blocks_treat_setext_marker_as_rule_not_heading():
    blocks = parse_markdown_blocks(
        [
            "**Datum:** [fyll i]",
            "**Plats:** [fyll i]",
            "---",
        ]
    )

    assert blocks == [
        DocumentBlock(
            kind="paragraph",
            text="Datum: [fyll i]\nPlats: [fyll i]",
            runs=(
                InlineTextRun("Datum:", bold=True),
                InlineTextRun(" [fyll i]"),
                InlineTextRun("\n"),
                InlineTextRun("Plats:", bold=True),
                InlineTextRun(" [fyll i]"),
            ),
        ),
        DocumentBlock(kind="empty"),
    ]


def test_markdown_blocks_preserve_inline_run_semantics():
    blocks = parse_markdown_blocks(["**Fet** och _kursiv_ samt `kod` och ~~struken~~"])

    paragraph = blocks[0]
    assert paragraph.text == "Fet och kursiv samt kod och struken"
    assert paragraph.runs == (
        InlineTextRun("Fet", bold=True),
        InlineTextRun(" och "),
        InlineTextRun("kursiv", italic=True),
        InlineTextRun(" samt "),
        InlineTextRun("kod", code=True),
        InlineTextRun(" och "),
        InlineTextRun("struken", strikethrough=True),
    )


# --- DOCX rendering ---


def test_render_docx_valid_blob():
    blob, mimetype, filename = _render_service.render_document("Hello world", "docx", step_order=1)
    assert isinstance(blob, bytes)
    assert len(blob) > 0
    # DOCX is a ZIP file — starts with PK magic bytes
    assert blob[:2] == b"PK"


def test_render_docx_correct_mime():
    _, mimetype, _ = _render_service.render_document("Test", "docx", step_order=1)
    assert (
        mimetype
        == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


def test_render_docx_filename_pattern():
    _, _, filename = _render_service.render_document("Test", "docx", step_order=5)
    assert filename == "step_5_output.docx"


def test_render_docx_empty_output_still_valid():
    """Empty markdown should still produce a readable DOCX file."""
    blob, _, _ = _render_service.render_document("", "docx", step_order=1)
    doc = Document(io.BytesIO(blob))
    assert isinstance(blob, bytes)
    assert len(doc.paragraphs) >= 1


def test_render_docx_preserves_swedish_characters():
    """Swedish characters should survive DOCX rendering."""
    text = "Svenska tecken: å ä ö"
    blob, _, _ = _render_service.render_document(text, "docx", step_order=1)
    doc = Document(io.BytesIO(blob))
    all_text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    assert "å" in all_text
    assert "ä" in all_text
    assert "ö" in all_text


def test_render_docx_markdown_table_creates_table():
    """Markdown table syntax should become a DOCX table."""
    text = "| Namn | Värde |\n| --- | --- |\n| Kommun | Sundsvall |"
    blob, _, _ = _render_service.render_document(text, "docx", step_order=1)
    doc = Document(io.BytesIO(blob))
    assert len(doc.tables) == 1
    assert doc.tables[0].cell(0, 0).text == "Namn"
    assert doc.tables[0].cell(1, 1).text == "Sundsvall"


def test_render_docx_markdown_inline_syntax_as_readable_text():
    text = (
        "**Medarbetare:** Leona\n\n"
        "---\n\n"
        "1) **Sammanfattning**\n"
        "Träna `skrivande` och ~~gamla mål~~."
    )

    blob, _, _ = _render_service.render_document(text, "docx", step_order=1)

    doc = Document(io.BytesIO(blob))
    all_text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    assert "Medarbetare: Leona" in all_text
    assert "Sammanfattning" in all_text
    assert "Träna skrivande och gamla mål." in all_text
    assert "**" not in all_text
    assert "---" not in all_text
    assert "`" not in all_text
    assert "~~" not in all_text


def test_render_docx_preserves_inline_bold_runs():
    blob, _, _ = _render_service.render_document("**Medarbetare:** Leona", "docx", step_order=1)

    doc = Document(io.BytesIO(blob))
    paragraph = doc.paragraphs[0]
    assert paragraph.runs[0].text == "Medarbetare:"
    assert paragraph.runs[0].bold is True
    assert "".join(run.text for run in paragraph.runs) == "Medarbetare: Leona"


def test_render_structured_docx_uses_schema_titles_and_tables():
    """Validated JSON contracts should render as semantic DOCX content."""
    blob, _, _ = _render_service.render_structured_document(
        {
            "summary": "Kort sammanfattning",
            "actions": [
                {"task": "Läsa mer", "owner": "Leona"},
                {"owner": "Lärare", "task": "Följa upp"},
            ],
        },
        "docx",
        step_order=1,
        schema={
            "title": "Utvecklingsrapport",
            "type": "object",
            "properties": {
                "summary": {"title": "Sammanfattning", "type": "string"},
                "actions": {
                    "title": "Åtgärder",
                    "type": "array",
                    "items": {
                        "type": "object",
                        "properties": {
                            "owner": {"title": "Ansvarig", "type": "string"},
                            "task": {"title": "Uppgift", "type": "string"},
                        },
                    },
                },
            },
        },
    )

    doc = Document(io.BytesIO(blob))
    all_text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    assert "Utvecklingsrapport" in all_text
    assert "Sammanfattning" in all_text
    assert "Kort sammanfattning" in all_text
    assert len(doc.tables) == 1
    assert doc.tables[0].cell(0, 0).text == "Ansvarig"
    assert doc.tables[0].cell(0, 1).text == "Uppgift"
    assert doc.tables[0].cell(1, 1).text == "Läsa mer"


def test_render_structured_docx_pins_null_and_empty_array_values():
    """Missing structured values should stay visible instead of disappearing."""
    blob, _, _ = _render_service.render_structured_document(
        {"summary": None, "notes": []},
        "docx",
        step_order=1,
        schema={
            "type": "object",
            "properties": {
                "summary": {"title": "Sammanfattning", "type": "string"},
                "notes": {"title": "Anteckningar", "type": "array"},
            },
        },
    )

    doc = Document(io.BytesIO(blob))
    all_text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    assert "Sammanfattning: -" in all_text
    assert "Anteckningar" in all_text
    assert "-" in all_text


def test_render_structured_docx_escapes_scalar_markdown_control_text():
    """JSON strings are data and must not become document structure."""
    blob, _, _ = _render_service.render_structured_document(
        {"summary": "Rad ett\n# Inte en rubrik\n```inte kod```"},
        "docx",
        step_order=1,
        schema={
            "type": "object",
            "properties": {
                "summary": {"title": "Sammanfattning", "type": "string"},
            },
        },
    )

    doc = Document(io.BytesIO(blob))
    paragraph_texts = [paragraph.text for paragraph in doc.paragraphs]
    assert (
        "Sammanfattning: Rad ett\n# Inte en rubrik\n'''inte kod'''" in paragraph_texts
    )
    assert "Inte en rubrik" not in paragraph_texts


def test_render_structured_docx_table_cells_preserve_pipe_characters():
    """Escaped Markdown table pipes should round-trip into DOCX table cells."""
    blob, _, _ = _render_service.render_structured_document(
        {"rows": [{"name": "A | B", "value": "C"}]},
        "docx",
        step_order=1,
        schema={"type": "object", "properties": {"rows": {"type": "array"}}},
    )

    doc = Document(io.BytesIO(blob))
    assert len(doc.tables) == 1
    assert doc.tables[0].cell(1, 0).text == "A | B"


def test_render_structured_pdf_table_does_not_truncate_long_cell_values():
    """PDF fallback tables should prefer readable wrapping over lossy truncation."""
    long_value = (
        "Detaljerad uppföljning med tillräckligt många ord för att behöva "
        "radbrytas i PDF-tabellen och ändå behålla unique-tail-token."
    )

    blob, _, _ = _render_service.render_structured_document(
        {"actions": [{"owner": "Leona", "task": long_value}]},
        "pdf",
        step_order=1,
        schema={"type": "object", "properties": {"actions": {"type": "array"}}},
    )

    with pdfplumber.open(io.BytesIO(blob)) as pdf:
        page_text = "\n".join(page.extract_text() or "" for page in pdf.pages)
    normalized_text = page_text.replace("-\n", "").replace("\n", " ")
    assert "unique" in normalized_text
    assert "tail" in normalized_text
    assert "token" in normalized_text


def test_render_structured_pdf_table_repeats_headers_after_page_break():
    rows = [
        {
            "owner": f"Person {index}",
            "task": (
                "Följ upp utvecklingsmål och dokumentera status med tillräcklig "
                f"detalj för radbrytning {index}."
            ),
        }
        for index in range(45)
    ]

    blob, _, _ = _render_service.render_structured_document(
        {"actions": rows},
        "pdf",
        step_order=1,
        schema={
            "type": "object",
            "properties": {
                "actions": {
                    "type": "array",
                    "items": {
                        "type": "object",
                        "properties": {
                            "owner": {"title": "Ansvarig", "type": "string"},
                            "task": {"title": "Uppgift", "type": "string"},
                        },
                    },
                }
            },
        },
    )

    with pdfplumber.open(io.BytesIO(blob)) as pdf:
        assert len(pdf.pages) > 1
        second_page_text = pdf.pages[1].extract_text() or ""
    assert "Ansvarig" in second_page_text
    assert "Uppgift" in second_page_text


def test_pdf_renderer_uses_safe_semantic_weasyprint_options():
    captured_html_kwargs: dict[str, object] = {}
    captured_pdf_kwargs: dict[str, object] = {}

    class _FakeHtml:
        def __init__(self, **kwargs: object) -> None:
            captured_html_kwargs.update(kwargs)

        def write_pdf(self, **kwargs: object) -> bytes:
            captured_pdf_kwargs.update(kwargs)
            return b"%PDF-stub"

    class _FakeCss:
        def __init__(self, **kwargs: object) -> None:
            pass

    class _FakeFontConfiguration:
        pass

    with patch(
        "intric.flows.runtime.document_rendering.weasyprint_renderer."
        "_load_weasyprint_api",
        return_value=(_FakeHtml, _FakeCss, _FakeFontConfiguration),
    ):
        rendered = WeasyPrintDocumentRenderer().render(
            [DocumentBlock(kind="paragraph", text="Tillgänglig rapport")],
            step_order=1,
        )

    assert rendered.blob == b"%PDF-stub"
    assert captured_html_kwargs["media_type"] == "print"
    assert captured_html_kwargs["url_fetcher"] is _deny_external_fetch
    assert captured_pdf_kwargs["pdf_tags"] is True
    assert captured_pdf_kwargs["srgb"] is True
    assert captured_pdf_kwargs["font_config"] is not None
    assert "pdf_variant" not in captured_pdf_kwargs


def test_html_document_escapes_untrusted_text():
    html = blocks_to_html_document(
        [DocumentBlock(kind="paragraph", text='<script>alert("x")</script>')],
        title='Output "test"',
    )

    assert "<script>" not in html
    assert "&lt;script&gt;" in html
    assert "Output &quot;test&quot;" in html


def test_markdown_image_url_is_rendered_as_text_not_external_resource():
    blocks = parse_markdown_blocks(["![Alt text](https://example.com/tracker.png)"])

    html = blocks_to_html_document(blocks, title="Output")

    assert "<img" not in html
    assert "https://example.com/tracker.png" not in html
    assert "<p>Alt text</p>" in html


def test_html_document_preserves_inline_semantics():
    html = blocks_to_html_document(
        [
            DocumentBlock(
                kind="paragraph",
                text="Viktig text",
                runs=(
                    InlineTextRun("Viktig", bold=True),
                    InlineTextRun(" text"),
                ),
            )
        ],
        title="Output",
    )

    assert "<strong>Viktig</strong> text" in html


def test_html_document_marks_table_headers_as_column_headers():
    html = blocks_to_html_document(
        [
            DocumentBlock(
                kind="table",
                rows=(("Område", "Status"), ("Svenska", "Utvecklas")),
            )
        ],
        title="Output",
    )

    assert '<th scope="col">Område</th>' in html
    assert '<th scope="col">Status</th>' in html


def test_html_document_keeps_header_semantics_for_header_only_tables():
    html = blocks_to_html_document(
        [DocumentBlock(kind="table", rows=(("Område", "Status"),))],
        title="Output",
    )

    assert "<thead>" in html
    assert '<th scope="col">Område</th>' in html
    assert "<td>Område</td>" not in html


def test_pdf_renderer_denies_external_resource_fetches():
    with pytest.raises(Exception, match="External PDF resource loading is disabled"):
        _deny_external_fetch("https://example.com/tracker.png")


def test_pdf_renderer_suppresses_noisy_font_subset_logs():
    configure_weasyprint_dependency_logging()

    import logging

    assert logging.getLogger("fontTools.subset").getEffectiveLevel() >= logging.WARNING
    assert logging.getLogger("weasyprint").getEffectiveLevel() >= logging.WARNING


def test_document_render_service_can_use_injected_renderer():
    class _FakeRenderer:
        output_type = "pdf"

        def render(
            self,
            blocks: Sequence[DocumentBlock],
            *,
            step_order: int,
        ) -> RenderedDocument:
            assert [(block.kind, block.text) for block in blocks] == [
                ("paragraph", "Hello")
            ]
            return RenderedDocument(
                blob=b"custom",
                mimetype="application/custom-pdf",
                filename=f"custom-{step_order}.pdf",
            )

    service = DocumentRenderService(renderers=(_FakeRenderer(),))

    assert service.render_document("Hello", "pdf", step_order=9) == (
        b"custom",
        "application/custom-pdf",
        "custom-9.pdf",
    )


def test_document_render_service_raises_stable_error_without_leaking_details():
    class _FailingRenderer:
        output_type = "pdf"

        def render(
            self,
            blocks: Sequence[DocumentBlock],
            *,
            step_order: int,
        ) -> RenderedDocument:
            raise RuntimeError("internal path /tmp/secret-file")

    service = DocumentRenderService(renderers=(_FailingRenderer(),))

    with pytest.raises(TypedIOValidationException) as exc_info:
        service.render_document("Hello", "pdf", step_order=1)

    assert exc_info.value.code == "typed_io_render_failed"
    assert str(exc_info.value) == "Document render failed."
    assert "secret-file" not in str(exc_info.value)


def test_document_render_service_rejects_outputs_over_render_limits():
    service = DocumentRenderService(renderers=(WeasyPrintDocumentRenderer(),))
    text = "x" * 500_001

    with pytest.raises(TypedIOValidationException, match="too large"):
        service.render_document(text, "pdf", step_order=1)


def test_document_render_service_uses_injected_render_limits():
    service = DocumentRenderService(
        renderers=(WeasyPrintDocumentRenderer(),),
        limits=DocumentRenderLimits(max_source_chars=5),
    )

    with pytest.raises(TypedIOValidationException) as exc_info:
        service.render_document("123456", "pdf", step_order=1)

    assert exc_info.value.context == {
        "metric": "source_chars",
        "actual": 6,
        "limit": 5,
    }


def test_document_render_service_rejects_too_many_table_cells():
    service = DocumentRenderService(renderers=(WeasyPrintDocumentRenderer(),))
    row = tuple(f"cell-{index}" for index in range(50))
    rows = tuple(row for _ in range(1_001))

    with pytest.raises(TypedIOValidationException) as exc_info:
        service.render_blocks(
            [DocumentBlock(kind="table", rows=rows)],
            "pdf",
            step_order=1,
        )

    assert exc_info.value.code == "typed_io_render_failed"
    assert exc_info.value.context == {
        "metric": "table_cells",
        "actual": 50_050,
        "limit": 50_000,
    }


def test_document_render_service_rejects_too_many_structured_list_items():
    data: list[object] = [None] * 5_001

    with pytest.raises(TypedIOValidationException) as exc_info:
        _render_service.render_structured_document(data, "pdf", step_order=1)

    assert exc_info.value.code == "typed_io_render_failed"
    assert exc_info.value.context == {
        "metric": "list_items",
        "actual": 5_001,
        "limit": 5_000,
    }


def test_document_render_service_rejects_deep_structured_values_before_conversion():
    data: list[object] = []
    current = data
    for _ in range(33):
        child: list[object] = []
        current.append(child)
        current = child

    with pytest.raises(TypedIOValidationException) as exc_info:
        _render_service.render_structured_document(data, "docx", step_order=1)

    assert exc_info.value.code == "typed_io_render_failed"
    assert exc_info.value.context is not None
    assert exc_info.value.context["metric"] == "structured_depth"


def test_render_docx_markdown_lists_and_code_blocks():
    """Lists and fenced code blocks should be represented in DOCX text."""
    text = "# Titel\n\n- punkt ett\n- punkt två\n\n```python\nprint('hej')\n```"
    blob, _, _ = _render_service.render_document(text, "docx", step_order=1)
    doc = Document(io.BytesIO(blob))
    all_text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    assert "Titel" in all_text
    assert "punkt ett" in all_text
    assert "print('hej')" in all_text


def test_render_docx_very_long_output():
    """Very long markdown output should produce a valid DOCX without exceptions."""
    text = ("Rad med innehåll och åäö.\n" * 5000).strip()
    blob, _, _ = _render_service.render_document(text, "docx", step_order=1)
    doc = Document(io.BytesIO(blob))
    assert len(blob) > 0
    assert any("åäö" in paragraph.text for paragraph in doc.paragraphs)


def test_render_docx_still_works_when_package_default_template_is_missing():
    """Renderer should not depend on python-docx package template layout."""
    with patch("docx.api._default_docx_path", return_value="/tmp/missing-default.docx"):
        blob, _, _ = _render_service.render_document("Fallback template test", "docx", step_order=1)
    doc = Document(io.BytesIO(blob))
    all_text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    assert "Fallback template test" in all_text


# --- Error handling ---


def test_render_unsupported_type_raises():
    with pytest.raises(TypedIOValidationException, match="Unsupported document type"):
        _render_service.render_document("Test", "html", step_order=1)


def test_render_unsupported_error_code():
    with pytest.raises(TypedIOValidationException) as exc_info:
        _render_service.render_document("Test", "html", step_order=1)
    assert exc_info.value.code == "typed_io_render_failed"


# --- Unicode rendering ---


def test_render_pdf_unicode_characters():
    """Em-dash, Swedish chars, curly quotes must render without error."""
    text = "Em-dash \u2014 and Swedish: \u00e5\u00e4\u00f6 and curly \u201cquotes\u201d"
    blob, mimetype, filename = _render_service.render_document(text, "pdf", step_order=1)
    assert isinstance(blob, bytes)
    assert len(blob) > 0
    assert blob[:5] == b"%PDF-"


def test_render_pdf_font_fallback():
    """System font fallback should keep Unicode PDF rendering available."""
    blob, _, _ = _render_service.render_document(
        "Em-dash \u2014 and curly \u201cquotes\u201d",
        "pdf",
        step_order=1,
    )
    assert isinstance(blob, bytes)
    assert blob[:5] == b"%PDF-"
