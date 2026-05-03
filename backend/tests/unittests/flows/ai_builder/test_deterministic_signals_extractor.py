"""Contract tests for `extract_deterministic_signals`.

Pins the dispatcher surface callers rely on:

- every call populates file-level signals (mime_type, extension,
  size_bytes) regardless of mime family,
- extension is inferred from the filename suffix first and falls back
  to the `MIMETYPE_EXTENSIONS_MAPPER` registry when the filename
  carries no informative suffix,
- the text extractor populates `bullet_density` from bytes it can
  decode as UTF-8 (with replacement on invalid sequences),
- unsupported mimes still return a valid `DeterministicSignals` with
  the default empty / None per-mime fields, so upstream callers never
  need a None check on the top-level return.
"""

from __future__ import annotations

from html import escape
from io import BytesIO

import numpy as np
import pytest
import soundfile
from docx import Document
from openpyxl import Workbook

from intric.flows.ai_builder.attachment_observation import DeterministicSignals
from intric.flows.ai_builder.deterministic_signals_extractor import (
    extract_deterministic_signals,
)

_DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
_PDF_MIME = "application/pdf"
_XLSX_MIME = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
_CSV_MIME = "text/csv"
_WAV_MIME = "audio/wav"
_OGG_MIME = "audio/ogg"


def _bytes(text: str) -> bytes:
    return text.encode("utf-8")


def _build_docx(
    *,
    headings: list[tuple[int, str]] | None = None,
    paragraphs: list[str] | None = None,
    tables: list[tuple[int, int]] | None = None,
) -> bytes:
    """Build a minimal in-memory DOCX containing the declared structure.

    Used by tests so the extractor's DOCX branch is exercised against
    real openxml bytes, not a hand-rolled fake; anything python-docx
    itself might change (namespace URIs, relationship IDs) is therefore
    covered by round-trip from the upstream writer to the extractor's
    reader.
    """
    document = Document()
    for level, text in headings or []:
        document.add_heading(text, level=level)
    for paragraph in paragraphs or []:
        document.add_paragraph(paragraph)
    for rows, cols in tables or []:
        document.add_table(rows=rows, cols=cols)
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


class TestFileLevelSignals:
    def test_dispatcher_sets_mime_extension_size_for_text(self):
        signals = extract_deterministic_signals(
            raw_bytes=_bytes("hello"),
            mime="text/plain",
            filename="notes.txt",
        )

        assert signals.mime_type == "text/plain"
        assert signals.extension == "txt"
        assert signals.size_bytes == 5

    def test_size_bytes_reflects_actual_byte_length(self):
        payload = b"\xef\xbb\xbfhello"
        signals = extract_deterministic_signals(
            raw_bytes=payload,
            mime="text/plain",
            filename="notes.txt",
        )

        assert signals.size_bytes == len(payload)

    def test_returns_deterministic_signals_instance(self):
        signals = extract_deterministic_signals(
            raw_bytes=b"",
            mime="application/pdf",
            filename="empty.pdf",
        )

        assert isinstance(signals, DeterministicSignals)


class TestExtensionInference:
    def test_uses_filename_suffix_lowercased_without_dot(self):
        signals = extract_deterministic_signals(
            raw_bytes=b"",
            mime="text/plain",
            filename="Report.TXT",
        )

        assert signals.extension == "txt"

    def test_falls_back_to_mime_mapping_when_filename_has_no_suffix(self):
        signals = extract_deterministic_signals(
            raw_bytes=b"",
            mime="application/pdf",
            filename="no-extension",
        )

        assert signals.extension == "pdf"

    def test_empty_extension_when_mime_and_filename_yield_nothing(self):
        signals = extract_deterministic_signals(
            raw_bytes=b"",
            mime="application/x-unknown",
            filename="blob",
        )

        assert signals.extension == ""

    def test_hidden_file_with_no_suffix_still_falls_back_to_mime(self):
        signals = extract_deterministic_signals(
            raw_bytes=b"",
            mime="text/plain",
            filename=".gitignore",
        )

        assert signals.extension == "txt"


class TestTextBulletDensity:
    def test_plain_text_without_bullets_has_zero_density(self):
        signals = extract_deterministic_signals(
            raw_bytes=_bytes("line one\nline two\nline three"),
            mime="text/plain",
            filename="notes.txt",
        )

        assert signals.bullet_density == 0.0

    def test_all_bullet_lines_has_density_one(self):
        signals = extract_deterministic_signals(
            raw_bytes=_bytes("- a\n- b\n- c"),
            mime="text/plain",
            filename="notes.txt",
        )

        assert signals.bullet_density == 1.0

    def test_mixed_bullet_and_prose_has_expected_ratio(self):
        signals = extract_deterministic_signals(
            raw_bytes=_bytes("intro\n* bullet one\n* bullet two\noutro"),
            mime="text/plain",
            filename="notes.txt",
        )

        assert signals.bullet_density is not None
        assert 0.49 <= signals.bullet_density <= 0.51

    def test_recognises_asterisk_dash_plus_and_typographic_bullets(self):
        content = "* a\n- b\n+ c\n• d\n· e"
        signals = extract_deterministic_signals(
            raw_bytes=_bytes(content),
            mime="text/plain",
            filename="notes.txt",
        )

        assert signals.bullet_density == 1.0

    def test_blank_lines_ignored_when_computing_density(self):
        signals = extract_deterministic_signals(
            raw_bytes=_bytes("\n- a\n\n- b\n\n"),
            mime="text/plain",
            filename="notes.txt",
        )

        assert signals.bullet_density == 1.0

    def test_empty_text_leaves_bullet_density_none(self):
        signals = extract_deterministic_signals(
            raw_bytes=b"",
            mime="text/plain",
            filename="notes.txt",
        )

        assert signals.bullet_density is None

    def test_indented_bullet_lines_still_count(self):
        signals = extract_deterministic_signals(
            raw_bytes=_bytes("  - indented\n    * also indented"),
            mime="text/plain",
            filename="notes.md",
        )

        assert signals.bullet_density == 1.0

    def test_invalid_utf8_bytes_do_not_crash(self):
        payload = b"valid line\n\xff\xfeinvalid\n- bullet"
        signals = extract_deterministic_signals(
            raw_bytes=payload,
            mime="text/plain",
            filename="notes.txt",
        )

        assert signals.bullet_density is not None


class TestMarkdownIsTreatedAsText:
    def test_markdown_mime_populates_bullet_density(self):
        signals = extract_deterministic_signals(
            raw_bytes=_bytes("- one\n- two"),
            mime="text/markdown",
            filename="README.md",
        )

        assert signals.mime_type == "text/markdown"
        assert signals.extension == "md"
        assert signals.bullet_density == 1.0


class TestUnsupportedMimeFallback:
    def test_unknown_mime_returns_file_level_signals_without_per_mime_fields(self):
        signals = extract_deterministic_signals(
            raw_bytes=b"%PDF-1.4 minimal",
            mime="application/pdf",
            filename="doc.pdf",
        )

        assert signals.mime_type == "application/pdf"
        assert signals.extension == "pdf"
        assert signals.size_bytes == len(b"%PDF-1.4 minimal")
        assert signals.page_count is None
        assert signals.table_count is None
        assert signals.form_fields == []
        assert signals.placeholder_tokens == []
        assert signals.bullet_density is None


class TestDocxExtractor:
    """DOCX attachments carry the highest-signal structure for AI Builder
    flow-archetype selection: a document with `{{placeholder}}` tokens
    is a `document_to_docx_template` candidate; one with rich heading
    structure and tables is a `document_to_structured_report` candidate.
    The extractor surfaces those signals mechanically before any LLM
    pass sees the file.
    """

    def test_headings_are_extracted_with_levels_and_text(self) -> None:
        raw = _build_docx(
            headings=[
                (1, "Contract"),
                (2, "Scope"),
                (2, "Deliverables"),
            ],
        )
        signals = extract_deterministic_signals(
            raw_bytes=raw,
            mime=_DOCX_MIME,
            filename="contract.docx",
        )

        assert signals.heading_tree is not None
        headings = [(h.level, h.text) for h in signals.heading_tree]
        assert headings == [
            (1, "Contract"),
            (2, "Scope"),
            (2, "Deliverables"),
        ]

    def test_placeholder_tokens_are_extracted_from_body_text(self) -> None:
        """Jinja-style `{{ var }}` placeholders are the dominant signal
        for `document_to_docx_template` — the planner uses the set
        of tokens to score whether the attachment is a fill-in
        template versus a generated output sample.
        """
        raw = _build_docx(
            paragraphs=[
                "Hello {{ customer_name }}, welcome to {{ company }}.",
                "Your order number is {{order_id}}.",
                "Regular prose with no tokens.",
                "{{customer_name}} again — duplicates collapse to a single token.",
            ],
        )
        signals = extract_deterministic_signals(
            raw_bytes=raw,
            mime=_DOCX_MIME,
            filename="template.docx",
        )

        assert set(signals.placeholder_tokens) == {
            "customer_name",
            "company",
            "order_id",
        }

    def test_table_count_and_dimensions_match_document(self) -> None:
        raw = _build_docx(tables=[(2, 3), (4, 5)])
        signals = extract_deterministic_signals(
            raw_bytes=raw,
            mime=_DOCX_MIME,
            filename="report.docx",
        )

        assert signals.table_count == 2
        dims = [(t.rows, t.cols) for t in signals.table_dimensions]
        assert dims == [(2, 3), (4, 5)]

    def test_docx_without_placeholders_returns_empty_token_list(self) -> None:
        raw = _build_docx(paragraphs=["Plain prose.", "No tokens here."])
        signals = extract_deterministic_signals(
            raw_bytes=raw,
            mime=_DOCX_MIME,
            filename="prose.docx",
        )
        assert signals.placeholder_tokens == []

    def test_empty_docx_returns_empty_structural_fields(self) -> None:
        raw = _build_docx()
        signals = extract_deterministic_signals(
            raw_bytes=raw,
            mime=_DOCX_MIME,
            filename="empty.docx",
        )

        assert signals.mime_type == _DOCX_MIME
        assert signals.extension == "docx"
        assert signals.heading_tree == []
        assert signals.placeholder_tokens == []
        assert signals.table_count == 0
        assert signals.table_dimensions == []

    def test_malformed_docx_bytes_fall_back_to_file_level_signals(self) -> None:
        """DOCX is a ZIP; non-ZIP bytes must not crash the extractor.
        Upstream MIME classification is permissive (cache key derives
        from every byte stream reaching the pipeline), so a malformed
        payload should still produce a valid DeterministicSignals.
        """
        signals = extract_deterministic_signals(
            raw_bytes=b"not a real docx file",
            mime=_DOCX_MIME,
            filename="broken.docx",
        )

        assert signals.mime_type == _DOCX_MIME
        assert signals.extension == "docx"
        assert signals.heading_tree is None
        assert signals.table_count is None
        assert signals.placeholder_tokens == []


def _build_pdf(*, pages: list[str]) -> bytes:
    """Build a minimal in-memory PDF with the given page texts.

    Empty-string page entries produce blank pages (no rendered text),
    which is how the tests exercise the scanned-PDF signal. WeasyPrint
    keeps PDF fixtures on the same generation stack as flow PDF outputs
    instead of retaining a second PDF writer only for tests.
    """
    try:
        from weasyprint import HTML
    except (ImportError, OSError) as exc:
        pytest.skip(f"WeasyPrint system libraries unavailable: {exc}")

    sections: list[str] = []
    for index, text in enumerate(pages):
        page_break = "page-break-after: always;" if index < len(pages) - 1 else ""
        content = f"<p>{escape(text)}</p>" if text else ""
        sections.append(
            f'<section style="min-height: 250mm; {page_break}">{content}</section>'
        )
    html = "<!doctype html><html><body>" + "".join(sections) + "</body></html>"
    return bytes(HTML(string=html).write_pdf())


class TestPdfExtractor:
    """PDF page_count and is_scanned_pdf are the structural signals the
    planner uses to decide whether to expect text-rich content or to
    offer an OCR / upload-retry cue — a scanned PDF with zero
    extractable text should not be asked to fill a text summary
    archetype without a user acknowledgement.
    """

    def test_page_count_matches_document(self) -> None:
        raw = _build_pdf(pages=["first", "second", "third"])
        signals = extract_deterministic_signals(
            raw_bytes=raw,
            mime=_PDF_MIME,
            filename="three-pages.pdf",
        )

        assert signals.page_count == 3

    def test_text_pdf_is_not_marked_scanned(self) -> None:
        raw = _build_pdf(pages=["Hello world, this is page one."])
        signals = extract_deterministic_signals(
            raw_bytes=raw,
            mime=_PDF_MIME,
            filename="textual.pdf",
        )

        assert signals.is_scanned_pdf is False

    def test_blank_pdf_is_marked_scanned(self) -> None:
        """A PDF with no extractable text is the proxy for image-only
        or scanned content — pdfplumber returns empty strings for
        every page. The extractor cannot distinguish a truly scanned
        PDF from a deliberately blank one without OCR, but either way
        the planner should treat the attachment as text-less.
        """
        raw = _build_pdf(pages=["", ""])
        signals = extract_deterministic_signals(
            raw_bytes=raw,
            mime=_PDF_MIME,
            filename="blank.pdf",
        )

        assert signals.is_scanned_pdf is True
        assert signals.page_count == 2

    def test_malformed_pdf_falls_back_to_file_level_signals(self) -> None:
        signals = extract_deterministic_signals(
            raw_bytes=b"not a real pdf",
            mime=_PDF_MIME,
            filename="broken.pdf",
        )

        assert signals.mime_type == _PDF_MIME
        assert signals.extension == "pdf"
        assert signals.page_count is None
        assert signals.is_scanned_pdf is None

    def test_pdf_extension_inferred_when_filename_has_no_suffix(self) -> None:
        raw = _build_pdf(pages=["page one text"])
        signals = extract_deterministic_signals(
            raw_bytes=raw,
            mime=_PDF_MIME,
            filename="no-extension",
        )

        assert signals.extension == "pdf"
        assert signals.page_count == 1


def _build_xlsx(*, rows: list[list[object]]) -> bytes:
    """Build a minimal in-memory XLSX with the given row values.

    Row 0 is treated as the header row by the extractor; subsequent
    rows are data. Cell types preserve through openpyxl: `str` stays
    string, `int` / `float` stays numeric, `datetime` stays datetime,
    `bool` stays boolean. Round-tripping through openpyxl's writer +
    reader ensures the extractor reads the same byte format the
    production pipeline would receive.
    """
    workbook = Workbook()
    sheet = workbook.active
    assert sheet is not None
    for row in rows:
        sheet.append(row)
    buffer = BytesIO()
    workbook.save(buffer)
    return buffer.getvalue()


class TestXlsxExtractor:
    """XLSX attachments are the dominant source for `extract_structured_fields`
    archetype routing. Column headers and per-column type inference let
    the planner decide whether to offer a JSON-schema extraction plan
    versus a text summary, without ever rendering the sheet.
    """

    def test_headers_extracted_from_first_row(self) -> None:
        raw = _build_xlsx(
            rows=[
                ["name", "age", "signup_date"],
                ["Alice", 30, "2024-01-01"],
            ]
        )
        signals = extract_deterministic_signals(
            raw_bytes=raw,
            mime=_XLSX_MIME,
            filename="users.xlsx",
        )

        assert signals.spreadsheet_headers == ["name", "age", "signup_date"]

    def test_row_count_excludes_header(self) -> None:
        raw = _build_xlsx(
            rows=[
                ["col"],
                ["a"],
                ["b"],
                ["c"],
            ]
        )
        signals = extract_deterministic_signals(
            raw_bytes=raw,
            mime=_XLSX_MIME,
            filename="three-rows.xlsx",
        )

        assert signals.row_count == 3

    def test_column_types_inferred_from_data_rows(self) -> None:
        raw = _build_xlsx(
            rows=[
                ["name", "age", "active"],
                ["Alice", 30, True],
                ["Bob", 25, False],
            ]
        )
        signals = extract_deterministic_signals(
            raw_bytes=raw,
            mime=_XLSX_MIME,
            filename="types.xlsx",
        )

        assert signals.spreadsheet_column_types == ["text", "numeric", "boolean"]

    def test_mixed_type_column_reports_mixed(self) -> None:
        """Columns with divergent cell types across data rows cannot be
        safely claimed for a single JSON-schema type — the planner
        needs to know so it doesn't emit a schema that rejects half
        the attachment's own data.
        """
        raw = _build_xlsx(
            rows=[
                ["value"],
                ["text here"],
                [42],
            ]
        )
        signals = extract_deterministic_signals(
            raw_bytes=raw,
            mime=_XLSX_MIME,
            filename="mixed.xlsx",
        )

        assert signals.spreadsheet_column_types == ["mixed"]

    def test_empty_column_reports_empty_type(self) -> None:
        raw = _build_xlsx(
            rows=[
                ["header_with_no_data"],
            ]
        )
        signals = extract_deterministic_signals(
            raw_bytes=raw,
            mime=_XLSX_MIME,
            filename="header-only.xlsx",
        )

        assert signals.spreadsheet_headers == ["header_with_no_data"]
        assert signals.row_count == 0
        assert signals.spreadsheet_column_types == ["empty"]

    def test_malformed_xlsx_falls_back_to_file_level_signals(self) -> None:
        signals = extract_deterministic_signals(
            raw_bytes=b"not a real xlsx file",
            mime=_XLSX_MIME,
            filename="broken.xlsx",
        )

        assert signals.mime_type == _XLSX_MIME
        assert signals.extension == "xlsx"
        assert signals.spreadsheet_headers is None
        assert signals.spreadsheet_column_types is None
        assert signals.row_count is None


class TestCsvExtractor:
    """CSV attachments surface the same structural signals as XLSX — the
    planner can't tell from file contents alone whether the user
    exported a spreadsheet or a database query result. Column types
    are inferred heuristically because every CSV value arrives as a
    string; the extractor tries numeric / boolean / date parsing
    before falling back to `"text"`.
    """

    def test_headers_from_first_line(self) -> None:
        raw = b"name,age,active\nAlice,30,true\nBob,25,false\n"
        signals = extract_deterministic_signals(
            raw_bytes=raw,
            mime=_CSV_MIME,
            filename="users.csv",
        )

        assert signals.spreadsheet_headers == ["name", "age", "active"]

    def test_row_count_excludes_header(self) -> None:
        raw = b"col\na\nb\nc\n"
        signals = extract_deterministic_signals(
            raw_bytes=raw,
            mime=_CSV_MIME,
            filename="three-rows.csv",
        )

        assert signals.row_count == 3

    def test_column_types_inferred_heuristically(self) -> None:
        raw = b"name,age,active,signup\nAlice,30,true,2024-01-01\nBob,25,false,2024-02-15\n"
        signals = extract_deterministic_signals(
            raw_bytes=raw,
            mime=_CSV_MIME,
            filename="typed.csv",
        )

        assert signals.spreadsheet_column_types == [
            "text",
            "numeric",
            "boolean",
            "date",
        ]

    def test_mixed_numeric_and_text_reports_mixed(self) -> None:
        raw = b"value\n42\nhello world\n"
        signals = extract_deterministic_signals(
            raw_bytes=raw,
            mime=_CSV_MIME,
            filename="mixed.csv",
        )

        assert signals.spreadsheet_column_types == ["mixed"]

    def test_empty_column_reports_empty_type(self) -> None:
        raw = b"only_header\n"
        signals = extract_deterministic_signals(
            raw_bytes=raw,
            mime=_CSV_MIME,
            filename="header-only.csv",
        )

        assert signals.spreadsheet_headers == ["only_header"]
        assert signals.row_count == 0
        assert signals.spreadsheet_column_types == ["empty"]

    def test_application_csv_mime_handled_same_as_text_csv(self) -> None:
        """Browsers send either `text/csv` or `application/csv` for the
        same file; both must reach the CSV extractor branch so the
        pipeline signal set is stable across browser quirks.
        """
        raw = b"x,y\n1,2\n"
        signals = extract_deterministic_signals(
            raw_bytes=raw,
            mime="application/csv",
            filename="nums.csv",
        )

        assert signals.spreadsheet_headers == ["x", "y"]
        assert signals.spreadsheet_column_types == ["numeric", "numeric"]

    def test_malformed_csv_bytes_do_not_crash(self) -> None:
        """CSV is permissive — stray quotes / unterminated fields pass
        through the stdlib parser as best-effort lines. The extractor
        should not crash on invalid UTF-8 either; `errors='replace'`
        produces best-effort cells rather than losing the file.
        """
        raw = b"col\n\xff\xfe garbled"
        signals = extract_deterministic_signals(
            raw_bytes=raw,
            mime=_CSV_MIME,
            filename="garbled.csv",
        )

        assert signals.mime_type == _CSV_MIME
        assert signals.extension == "csv"
        # Header round-trips as-is; garbled data row still counted
        assert signals.spreadsheet_headers == ["col"]
        assert signals.row_count == 1


def _build_wav(
    *, duration_seconds: float, channels: int, samplerate: int = 44100
) -> bytes:
    """Build an in-memory WAV with the given duration and channel count.

    Round-tripping through soundfile's writer + reader means the test
    exercises the same decoder the extractor uses in production.
    """
    sample_count = int(duration_seconds * samplerate)
    shape = (sample_count, channels) if channels > 1 else (sample_count,)
    data = np.zeros(shape, dtype="float32")
    buffer = BytesIO()
    soundfile.write(buffer, data, samplerate, format="WAV")
    return buffer.getvalue()


def _build_ogg(*, duration_seconds: float, samplerate: int = 44100) -> bytes:
    sample_count = int(duration_seconds * samplerate)
    data = np.zeros(sample_count, dtype="float32")
    buffer = BytesIO()
    soundfile.write(buffer, data, samplerate, format="OGG", subtype="VORBIS")
    return buffer.getvalue()


class TestAudioExtractor:
    """Audio attachments feed the `audio_transcription` capability. The
    extractor reads duration and channel count without decoding the
    full stream so the planner can size the transcription budget
    before any LLM call. Language detection requires transcription
    and stays out of scope here.
    """

    def test_duration_in_seconds_matches_wav_header(self) -> None:
        raw = _build_wav(duration_seconds=2.5, channels=1)
        signals = extract_deterministic_signals(
            raw_bytes=raw,
            mime=_WAV_MIME,
            filename="clip.wav",
        )

        assert signals.duration_seconds is not None
        assert 2.4 <= signals.duration_seconds <= 2.6

    def test_channel_count_reflects_mono_vs_stereo(self) -> None:
        mono = _build_wav(duration_seconds=0.5, channels=1)
        stereo = _build_wav(duration_seconds=0.5, channels=2)
        mono_signals = extract_deterministic_signals(
            raw_bytes=mono,
            mime=_WAV_MIME,
            filename="mono.wav",
        )
        stereo_signals = extract_deterministic_signals(
            raw_bytes=stereo,
            mime=_WAV_MIME,
            filename="stereo.wav",
        )

        assert mono_signals.channel_count == 1
        assert stereo_signals.channel_count == 2

    def test_ogg_vorbis_duration_extracted(self) -> None:
        """OGG Vorbis is one of the registered audio MIMEs; the extractor
        must handle it without a WAV-specific code path.
        """
        raw = _build_ogg(duration_seconds=1.0)
        signals = extract_deterministic_signals(
            raw_bytes=raw,
            mime=_OGG_MIME,
            filename="clip.ogg",
        )

        assert signals.duration_seconds is not None
        assert 0.9 <= signals.duration_seconds <= 1.1
        assert signals.channel_count == 1

    def test_language_hint_is_not_populated(self) -> None:
        """Language detection requires transcription — out of scope for
        the deterministic extractor. The field must stay None so the
        planner doesn't confuse "we haven't checked" with "we know
        it's null".
        """
        raw = _build_wav(duration_seconds=0.5, channels=1)
        signals = extract_deterministic_signals(
            raw_bytes=raw,
            mime=_WAV_MIME,
            filename="clip.wav",
        )

        assert signals.language_hint is None

    def test_malformed_audio_bytes_fall_back_to_file_level_signals(self) -> None:
        signals = extract_deterministic_signals(
            raw_bytes=b"not a real wav file",
            mime=_WAV_MIME,
            filename="broken.wav",
        )

        assert signals.mime_type == _WAV_MIME
        assert signals.extension == "wav"
        assert signals.duration_seconds is None
        assert signals.channel_count is None

    def test_unsupported_audio_format_falls_back_without_crash(self) -> None:
        """MP4 / WEBM containers require ffmpeg and are not covered by
        soundfile. The extractor should not crash — file-level signals
        still populate and the missing duration is the cue to upstream
        cache logic that deeper inspection is needed elsewhere.
        """
        signals = extract_deterministic_signals(
            raw_bytes=b"not a decodable m4a payload",
            mime="audio/mp4",
            filename="clip.m4a",
        )

        assert signals.mime_type == "audio/mp4"
        assert signals.duration_seconds is None
        assert signals.channel_count is None
