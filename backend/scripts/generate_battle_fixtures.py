#!/usr/bin/env python3
"""Generate deterministic local fixtures for the AI Builder battle corpus.

The script never calls a live API. Run it from the backend virtual environment:

    PYTHONPATH=src .venv/bin/python scripts/generate_battle_fixtures.py

It writes the fixture files and ``manifest.json`` below
``scripts/fixtures/ai_builder_battle``. The manifest pins each fixture's content
SHA-256; a battle case attaches a fixture by name and the harness uploads it
itself, so a fresh environment can run the suite without hand-provisioned file
IDs. Nothing here is environment-specific and nothing needs post-upload capture.

The checked-in ``01_protokoll_bun_2026_02_25.pdf`` is the canonical authentic
§17 protokollsutdrag (three pages) from the source and its SHA-256 is pinned
below. ``--source`` may initialize or replace that file only when the supplied
bytes match the pinned hash.
"""

from __future__ import annotations

import argparse
import hashlib
import io
import json
import shutil
import zipfile
from collections.abc import Mapping, Sequence
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

FIXTURE_DIR = Path(__file__).with_name("fixtures") / "ai_builder_battle"
MANIFEST_PATH = FIXTURE_DIR / "manifest.json"
MANIFEST_VERSION = 1
AUTHENTIC_PROTOCOL_SHA256 = (
    "cdab0e4471ca518fa5c7334a185a0c77da1c5743a49468e54d3ff094f824c6d8"
)
FIXED_TIMESTAMP = datetime(2026, 1, 1, 0, 0, 0)
FIXED_ZIP_TIMESTAMP = (2026, 1, 1, 0, 0, 0)


MATTER_FIXTURE_NAMES = (
    "01_protokoll_bun_2026_02_25.pdf",
    "02_tjansteskrivelse_underlag.docx",
    "03_barnkonsekvensanalys.docx",
    "04_remissvar.docx",
    "05_lokalkalkyl.csv",
    "06_tidigare_beslut.pdf",
)

GENERATED_FIXTURE_NAMES = (
    *MATTER_FIXTURE_NAMES,
    "decision_letter_template.docx",
    "example_report.docx",
    "generic_case_template.docx",
    "tjansteskrivelse_template.docx",
)


def _configure_document(document: Document) -> None:
    properties = document.core_properties
    properties.author = "Eneo battle corpus"
    properties.created = FIXED_TIMESTAMP
    properties.modified = FIXED_TIMESTAMP
    properties.last_printed = FIXED_TIMESTAMP
    properties.title = "Kommunalt ärendeunderlag"

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10)
    for name, size, bold in (
        ("Municipal Title", 20, True),
        ("Municipal Heading", 13, True),
        ("Municipal Metadata", 9, False),
    ):
        if name not in styles:
            style = styles.add_style(name, 1)
        else:
            style = styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = bold

    for section in document.sections:
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.4)
        section.right_margin = Cm(2.0)
        header = section.header.paragraphs[0]
        header.text = "SUNDSVALLS KOMMUN  |  BARN- OCH UTBILDNINGSFÖRVALTNINGEN"
        header.style = styles["Municipal Metadata"]
        footer = section.footer.paragraphs[0]
        footer.text = "Sundsvalls kommun"
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.style = styles["Municipal Metadata"]


def _document_bytes(document: Document) -> bytes:
    raw = io.BytesIO()
    document.save(raw)
    source = zipfile.ZipFile(io.BytesIO(raw.getvalue()))
    normalized = io.BytesIO()
    with (
        source,
        zipfile.ZipFile(
            normalized,
            mode="w",
            compression=zipfile.ZIP_DEFLATED,
            compresslevel=9,
        ) as target,
    ):
        for member in sorted(source.infolist(), key=lambda item: item.filename):
            info = zipfile.ZipInfo(member.filename, FIXED_ZIP_TIMESTAMP)
            info.compress_type = zipfile.ZIP_DEFLATED
            info.create_system = 3
            info.external_attr = 0o600 << 16
            target.writestr(info, source.read(member.filename))
    return normalized.getvalue()


def _write_docx(path: Path, build: object) -> None:
    if not callable(build):
        raise TypeError("DOCX builder must be callable")
    document = Document()
    _configure_document(document)
    build(document)
    path.write_bytes(_document_bytes(document))


def _add_metadata_table(document: Document, rows: Sequence[tuple[str, str]]) -> None:
    table = document.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value


def _build_decision_letter(document: Document) -> None:
    document.add_paragraph("BESLUTSBREV", style="Municipal Title")
    _add_metadata_table(
        document,
        (
            ("Diarienummer", "{{ diarienummer }}"),
            ("Beslutsdatum", "{{ beslutsdatum }}"),
            ("Handläggare", "{{ handlaggare }}"),
        ),
    )
    for heading, placeholder in (
        ("Beslut", "{{ beslut }}"),
        ("Motivering", "{{ motivering }}"),
        ("Villkor", "{{ villkor }}"),
        ("Så överklagar du", "{{ overklagandehanvisning }}"),
    ):
        document.add_paragraph(heading, style="Municipal Heading")
        document.add_paragraph(placeholder)


def _build_generic_template(document: Document) -> None:
    document.add_paragraph("ÄRENDERAPPORT", style="Municipal Title")
    _add_metadata_table(
        document,
        (("Kundnamn", "{{ kundnamn }}"), ("Ärende-ID", "{{ case_id }}")),
    )
    document.add_paragraph("Sammanfattning", style="Municipal Heading")
    document.add_paragraph("{{ mallinnehall }}")


def _build_tjansteskrivelse_template(document: Document) -> None:
    document.add_paragraph("TJÄNSTESKRIVELSE", style="Municipal Title")
    _add_metadata_table(
        document,
        (
            ("Diarienummer", "{{ diarienummer }}"),
            ("Handläggare", "{{ handlaggare }}"),
            ("Förvaltning", "{{ forvaltning }}"),
            ("Nämnd", "{{ namnd }}"),
            ("Beslutsdatum", "{{ beslutsdatum }}"),
        ),
    )
    for heading, placeholder in (
        ("Ärendet", "{{ sections.ärendet.text }}"),
        ("Bakgrund", "{{ sections.bakgrund.text }}"),
        ("Bedömning", "{{ sections.bedömning.text }}"),
        ("Konsekvenser", "{{ sections.konsekvenser.text }}"),
        (
            "Förslag till beslut",
            "{{ sections.förslag_till_beslut.text }}",
        ),
    ):
        document.add_paragraph(heading, style="Municipal Heading")
        document.add_paragraph(placeholder)


def _build_example_report(document: Document) -> None:
    document.add_paragraph("EXEMPELRAPPORT — KÄLLGENOMGÅNG", style="Municipal Title")
    _add_metadata_table(
        document,
        (("Rapportdatum", "2026-01-15"), ("Status", "Exempel, anonymiserat")),
    )
    for heading, body in (
        (
            "Källa 1 — Ärendeunderlag",
            "Redovisa dokumenttyp, datum, avsändare, relevanta fakta och osäkerheter.",
        ),
        (
            "Källa 2 — Tidigare beslut",
            "Återge beslutet separat och behåll den synliga källhänvisningen.",
        ),
        (
            "Samlad bedömning",
            "Syntetisera endast belagda uppgifter och visa motsägelser och luckor.",
        ),
    ):
        document.add_paragraph(heading, style="Municipal Heading")
        document.add_paragraph(body)


def _build_underlag(document: Document) -> None:
    document.add_paragraph(
        "Underlag till tjänsteskrivelse — Ny förskolestruktur i Njurunda",
        style="Municipal Title",
    )
    _add_metadata_table(
        document,
        (
            ("Diarienummer", "BUN-2026-00037-1"),
            ("Dokumentdatum", "2026-02-04"),
            ("Dokumenttyp", "Tjänsteskrivelseunderlag"),
        ),
    )
    document.add_paragraph("Ärende", style="Municipal Heading")
    document.add_paragraph(
        "Barnantalet i Njurunda har minskat. Förvaltningen föreslår att "
        "verksamheten vid Klockarbergets förskola i Kvissleby avvecklas."
    )
    document.add_paragraph("Bedömningspunkter", style="Municipal Heading")
    document.add_paragraph(
        "Barn och pedagoger ska erbjudas en planerad övergång. Personaltätheten "
        "ska efter förflyttningen i genomsnitt vara högst 5,0 barn per anställd. "
        "Lämpliga barngruppsstorlekar och en mångfald av förskolor ska bevaras."
    )


def _build_child_impact(document: Document) -> None:
    document.add_paragraph(
        "Barnkonsekvensanalys — Klockarbergets förskola",
        style="Municipal Title",
    )
    _add_metadata_table(
        document,
        (
            ("Diarienummer", "BUN-2026-00037-1"),
            ("Dokumentdatum", "2026-01-28"),
            ("Dokumenttyp", "Barnkonsekvensanalys"),
        ),
    )
    document.add_paragraph("Belagda konsekvenser", style="Municipal Heading")
    document.add_paragraph(
        "Avvecklingen innebär byte av förskola för berörda barn. Kontinuitet "
        "stärks om personal kan följa med barnen när lag, avtal och kommunens "
        "rutiner för personalförflyttning medger det. Närhet, trygghet och "
        "lämpliga barngrupper behöver följas upp under övergången."
    )
    document.add_paragraph("Uppföljning", style="Municipal Heading")
    document.add_paragraph(
        "Förvaltningen följer personaltäthet, barngruppsstorlek och trygghet före, "
        "under och efter övergången samt återrapporterar avvikelser till nämnden."
    )


def _build_consultation_response(document: Document) -> None:
    document.add_paragraph(
        "Sammanställt remissvar — Förskolestruktur Njurunda",
        style="Municipal Title",
    )
    _add_metadata_table(
        document,
        (
            ("Diarienummer", "BUN-2026-00037-1"),
            ("Dokumentdatum", "2026-02-18"),
            ("Dokumenttyp", "Remissvar"),
        ),
    )
    document.add_paragraph("Synpunkter", style="Municipal Heading")
    document.add_paragraph(
        "Synpunkterna betonar trygg övergång, information till vårdnadshavare "
        "och fortsatt variation mellan mindre och större förskolor."
    )
    document.add_paragraph("Datumuppgift att kontrollera", style="Municipal Heading")
    document.add_paragraph(
        "Remissammanställningen anger nämndens beslutsdatum som 2026-02-24."
    )


def _pdf_bytes(lines: Sequence[str]) -> bytes:
    def pdf_string(value: str) -> bytes:
        encoded = value.encode("cp1252")
        return (
            encoded.replace(b"\\", b"\\\\").replace(b"(", b"\\(").replace(b")", b"\\)")
        )

    commands = [b"BT", b"/F1 10 Tf", b"50 790 Td", b"13 TL"]
    for line in lines:
        commands.extend((b"(" + pdf_string(line) + b") Tj", b"T*"))
    commands.append(b"ET")
    stream = b"\n".join(commands) + b"\n"
    objects = (
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        (
            b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 595 842] "
            b"/Resources << /Font << /F1 4 0 R >> >> /Contents 5 0 R >>"
        ),
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica /Encoding /WinAnsiEncoding >>",
        b"<< /Length "
        + str(len(stream)).encode("ascii")
        + b" >>\nstream\n"
        + stream
        + b"endstream",
    )
    output = bytearray(b"%PDF-1.4\n%\xe2\xe3\xcf\xd3\n")
    offsets = [0]
    for index, value in enumerate(objects, start=1):
        offsets.append(len(output))
        output.extend(f"{index} 0 obj\n".encode("ascii"))
        output.extend(value)
        output.extend(b"\nendobj\n")
    xref_offset = len(output)
    output.extend(f"xref\n0 {len(objects) + 1}\n".encode("ascii"))
    output.extend(b"0000000000 65535 f \n")
    for offset in offsets[1:]:
        output.extend(f"{offset:010d} 00000 n \n".encode("ascii"))
    output.extend(
        (
            f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\n"
            f"startxref\n{xref_offset}\n%%EOF\n"
        ).encode("ascii")
    )
    return bytes(output)


def _sha256(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def _verify_authentic_protocol(path: Path) -> None:
    if not path.is_file():
        raise FileNotFoundError(f"Authentic protocol fixture is missing: {path}")
    actual_sha256 = _sha256(path)
    if actual_sha256 != AUTHENTIC_PROTOCOL_SHA256:
        raise ValueError(
            "Authentic protocol SHA-256 mismatch: "
            f"expected {AUTHENTIC_PROTOCOL_SHA256}, got {actual_sha256} for {path}"
        )


def _write_fixtures(*, source_path: Path | None = None) -> dict[str, str]:
    FIXTURE_DIR.mkdir(parents=True, exist_ok=True)
    canonical_protocol = FIXTURE_DIR / MATTER_FIXTURE_NAMES[0]
    if source_path is not None:
        _verify_authentic_protocol(source_path)
        if source_path.resolve() != canonical_protocol.resolve():
            shutil.copyfile(source_path, canonical_protocol)
    _verify_authentic_protocol(canonical_protocol)
    _write_docx(FIXTURE_DIR / "decision_letter_template.docx", _build_decision_letter)
    _write_docx(FIXTURE_DIR / "example_report.docx", _build_example_report)
    _write_docx(FIXTURE_DIR / "generic_case_template.docx", _build_generic_template)
    _write_docx(
        FIXTURE_DIR / "tjansteskrivelse_template.docx",
        _build_tjansteskrivelse_template,
    )
    _write_docx(FIXTURE_DIR / MATTER_FIXTURE_NAMES[1], _build_underlag)
    _write_docx(FIXTURE_DIR / MATTER_FIXTURE_NAMES[2], _build_child_impact)
    _write_docx(FIXTURE_DIR / MATTER_FIXTURE_NAMES[3], _build_consultation_response)
    (FIXTURE_DIR / MATTER_FIXTURE_NAMES[4]).write_text(
        "diarienummer,dokumentdatum,dokumenttyp,post,belopp_sek,kalla,status\n"
        "BUN-2026-00037-1,2026-02-03,Lokalkalkyl,avvecklingsarbete,180000,"
        "preliminar_kalkyl,beraknat\n"
        "BUN-2026-00037-1,2026-02-03,Lokalkalkyl,anpassning_mottagande_lokaler,"
        "420000,preliminar_kalkyl,beraknat\n"
        "BUN-2026-00037-1,2026-02-03,Lokalkalkyl,overgangsstod,,saknas,"
        "finansieringskalla_ej_faststalld\n",
        encoding="utf-8",
        newline="\n",
    )
    (FIXTURE_DIR / MATTER_FIXTURE_NAMES[5]).write_bytes(
        _pdf_bytes(
            (
                "Protokollsutdrag - Barn- och utbildningsnamndens arbetsutskott",
                "Diarienummer: BUN-2026-00037-1",
                "Sammantradesdatum: 2026-02-11",
                "Paragraf 11: Ny forskolestruktur i Njurunda",
                "Arbetsutskottet foreslar avveckling av Klockarbergets forskola.",
                "Forslaget motiveras av minskat barnantal i Njurunda.",
                "Beslutet ska forenas med trygg overgangen for barn och personal.",
                "Personaltatheten ska i genomsnitt vara hogst 5,0 barn per anstalld.",
                "Arkiveringsstampel:",
                "B",
                "E",
                "S",
                "L",
                "U",
                "T",
                "A",
                "D",
            )
        )
    )
    return {
        fixture_name: _sha256(FIXTURE_DIR / fixture_name)
        for fixture_name in sorted(GENERATED_FIXTURE_NAMES)
    }


def _write_manifest(content_sha256s: Mapping[str, str]) -> None:
    """Pin every fixture's bytes so a case can name one and mean exactly it.

    File IDs are deliberately absent: they are per-environment and per-upload,
    and hand-carrying them is what let the flagship runtime sentinel skip
    itself for six full suite runs. The harness uploads these bytes and derives
    the IDs it needs.
    """

    MANIFEST_PATH.write_text(
        json.dumps(
            {
                "version": MANIFEST_VERSION,
                "description": (
                    "Deterministic AI Builder battle fixtures. Regenerate with "
                    "backend/scripts/generate_battle_fixtures.py."
                ),
                "fixtures": dict(sorted(content_sha256s.items())),
            },
            ensure_ascii=False,
            indent=2,
        )
        + "\n",
        encoding="utf-8",
        newline="\n",
    )


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument(
        "--source",
        type=Path,
        help=(
            "Initialize the canonical protocol from a local file whose SHA-256 "
            "matches the pinned authentic source."
        ),
    )
    args = parser.parse_args()

    content_sha256s = _write_fixtures(source_path=args.source)
    _write_manifest(content_sha256s)

    print(f"fixture_directory={FIXTURE_DIR.resolve()}")
    for fixture_name, sha256 in sorted(content_sha256s.items()):
        print(f"{sha256}  {fixture_name}")
    print(f"manifest={MANIFEST_PATH.resolve()}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
