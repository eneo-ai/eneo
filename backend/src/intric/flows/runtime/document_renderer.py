"""PDF and DOCX generation for flow step outputs."""

from __future__ import annotations

import os
import re
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Literal, cast

from intric.main.exceptions import TypedIOValidationException

_PDF_FONT_ENV_VAR = "ENEO_FLOW_PDF_FONT_PATH"
# The production backend image installs fonts-dejavu-core. Custom images can
# point this env var at another TTF/OTF font with broader glyph coverage.
_DEFAULT_PDF_FONT_PATH = Path("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf")
_DOCX_TEMPLATE_PATH = Path(__file__).with_name("templates") / "default.docx"
_MARKDOWN_TABLE_SEPARATOR = re.compile(
    r"^\s*\|?\s*:?-{3,}:?\s*(\|\s*:?-{3,}:?\s*)+\|?\s*$"
)
_HEADING_PATTERN = re.compile(r"^(#{1,6})\s+(.+)$")
_BULLET_PATTERN = re.compile(r"^[-*]\s+(.+)$")
_NUMBERED_PATTERN = re.compile(r"^\d+\.\s+(.+)$")
_CORE_PDF_FONT_REPLACEMENTS = str.maketrans(
    {
        "\u2013": "-",
        "\u2014": "-",
        "\u2018": "'",
        "\u2019": "'",
        "\u201c": '"',
        "\u201d": '"',
        "\u2026": "...",
    }
)
_EMPTY_VALUE_PLACEHOLDER = "-"
_DocumentBlockKind = Literal[
    "empty",
    "heading",
    "paragraph",
    "bullet_list",
    "numbered_list",
    "code",
    "table",
]


@dataclass(frozen=True)
class _DocumentBlock:
    kind: _DocumentBlockKind
    text: str = ""
    level: int = 0
    items: tuple[str, ...] = ()
    rows: tuple[tuple[str, ...], ...] = ()


def render_document(
    text: str, output_type: str, *, step_order: int
) -> tuple[bytes, str, str]:
    """Render text as PDF or DOCX. Returns (blob, mimetype, filename)."""
    return _render_blocks(
        _parse_markdown_blocks(text.splitlines()),
        output_type,
        step_order=step_order,
    )


def render_structured_document(
    data: dict[str, Any] | list[Any],
    output_type: str,
    *,
    step_order: int,
    schema: dict[str, Any] | None = None,
) -> tuple[bytes, str, str]:
    """Render validated JSON data as a human-readable generated document."""
    return _render_blocks(
        _structured_data_to_blocks(data, schema=schema),
        output_type,
        step_order=step_order,
    )


def _render_blocks(
    blocks: list[_DocumentBlock],
    output_type: str,
    *,
    step_order: int,
) -> tuple[bytes, str, str]:
    renderers = {"pdf": _render_pdf, "docx": _render_docx}
    renderer = renderers.get(output_type)
    if renderer is None:
        raise TypedIOValidationException(
            f"Unsupported document type: {output_type}",
            code="typed_io_render_failed",
        )
    try:
        return renderer(blocks, step_order=step_order)
    except TypedIOValidationException:
        raise
    except Exception as exc:
        raise TypedIOValidationException(
            f"Document render failed: {exc}",
            code="typed_io_render_failed",
        ) from exc


def _render_pdf(
    blocks: list[_DocumentBlock], *, step_order: int
) -> tuple[bytes, str, str]:
    from fpdf import FPDF

    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    unicode_font_path = _resolved_pdf_unicode_font()
    if unicode_font_path is not None:
        pdf.add_font("EneoUnicode", fname=unicode_font_path)
        font_family = "EneoUnicode"
        uses_core_font = False
    else:
        font_family = "Helvetica"
        uses_core_font = True
    _write_pdf_blocks(
        pdf=pdf,
        blocks=blocks,
        font_family=font_family,
        uses_core_font=uses_core_font,
    )
    return bytes(pdf.output()), "application/pdf", f"step_{step_order}_output.pdf"


def _resolved_pdf_unicode_font() -> str | None:
    configured_path = os.environ.get(_PDF_FONT_ENV_VAR)
    if configured_path and Path(configured_path).is_file():
        return configured_path

    if _DEFAULT_PDF_FONT_PATH.is_file():
        return str(_DEFAULT_PDF_FONT_PATH)
    return None


def _text_for_core_pdf_font(text: str) -> str:
    return (
        text.translate(_CORE_PDF_FONT_REPLACEMENTS)
        .encode(
            "latin-1",
            errors="replace",
        )
        .decode("latin-1")
    )


def _structured_data_to_blocks(
    data: dict[str, Any] | list[Any],
    *,
    schema: dict[str, Any] | None,
) -> list[_DocumentBlock]:
    blocks: list[_DocumentBlock] = []
    schema_title = schema.get("title") if isinstance(schema, dict) else None
    if isinstance(schema_title, str) and schema_title.strip():
        blocks.append(
            _DocumentBlock(
                kind="heading",
                text=_single_line_text(schema_title),
                level=1,
            )
        )
        blocks.append(_DocumentBlock(kind="empty"))
    blocks.extend(
        _blocks_for_value(
            value=data,
            label=None,
            schema=schema if isinstance(schema, dict) else None,
            level=2 if blocks else 1,
        )
    )
    return _trim_trailing_empty_blocks(blocks)


def _blocks_for_value(
    *,
    value: Any,
    label: str | None,
    schema: dict[str, Any] | None,
    level: int,
) -> list[_DocumentBlock]:
    if isinstance(value, dict):
        return _blocks_for_object(
            value=cast(dict[str, Any], value),
            label=label,
            schema=schema,
            level=level,
        )
    if isinstance(value, list):
        return _blocks_for_array(
            value=cast(list[Any], value),
            label=label,
            schema=schema,
            level=level,
        )
    rendered_value = _scalar_text(value)
    if label is None:
        return [_DocumentBlock(kind="paragraph", text=rendered_value)]
    return [
        _DocumentBlock(kind="paragraph", text=f"{label}: {rendered_value}"),
        _DocumentBlock(kind="empty"),
    ]


def _blocks_for_object(
    *,
    value: dict[str, Any],
    label: str | None,
    schema: dict[str, Any] | None,
    level: int,
) -> list[_DocumentBlock]:
    blocks: list[_DocumentBlock] = []
    if label:
        blocks.append(_DocumentBlock(kind="heading", text=label, level=min(level, 4)))
        blocks.append(_DocumentBlock(kind="empty"))
        level += 1

    properties = _as_str_dict(schema.get("properties") if schema is not None else None)
    ordered_keys = _ordered_object_keys(value, properties)
    if not ordered_keys:
        blocks.append(_DocumentBlock(kind="paragraph", text=_EMPTY_VALUE_PLACEHOLDER))
        blocks.append(_DocumentBlock(kind="empty"))
        return blocks

    for key in ordered_keys:
        field_schema = _as_str_dict(properties.get(key) if properties else None)
        field_label = _schema_label(key, field_schema)
        blocks.extend(
            _blocks_for_value(
                value=value.get(key),
                label=field_label,
                schema=field_schema,
                level=level,
            )
        )
    return blocks


def _blocks_for_array(
    *,
    value: list[Any],
    label: str | None,
    schema: dict[str, Any] | None,
    level: int,
) -> list[_DocumentBlock]:
    blocks: list[_DocumentBlock] = []
    if not value:
        if label:
            blocks.append(
                _DocumentBlock(kind="heading", text=label, level=min(level, 4))
            )
            blocks.append(_DocumentBlock(kind="empty"))
        blocks.append(_DocumentBlock(kind="paragraph", text=_EMPTY_VALUE_PLACEHOLDER))
        blocks.append(_DocumentBlock(kind="empty"))
        return blocks

    if label:
        blocks.append(_DocumentBlock(kind="heading", text=label, level=min(level, 4)))
        blocks.append(_DocumentBlock(kind="empty"))

    if all(_is_scalar(item) for item in value):
        blocks.append(
            _DocumentBlock(
                kind="bullet_list",
                items=tuple(_scalar_text(item) for item in value),
            )
        )
        blocks.append(_DocumentBlock(kind="empty"))
        return blocks

    if _can_render_object_table(value):
        blocks.append(
            _table_block_for_objects(
                rows=cast(list[dict[str, Any]], value),
                schema=schema,
            )
        )
        blocks.append(_DocumentBlock(kind="empty"))
        return blocks

    item_schema = _as_str_dict(schema.get("items") if schema is not None else None)
    for index, item in enumerate(value, start=1):
        item_label = f"Item {index}"
        blocks.extend(
            _blocks_for_value(
                value=item,
                label=item_label,
                schema=item_schema,
                level=level + 1,
            )
        )
    return blocks


def _ordered_object_keys(
    value: dict[str, Any], properties: dict[str, Any] | None
) -> list[str]:
    ordered: list[str] = []
    if properties is not None:
        ordered.extend(key for key in properties if key in value)
    ordered.extend(key for key in value if key not in ordered)
    return ordered


def _schema_label(key: str, schema: dict[str, Any] | None) -> str:
    title = schema.get("title") if isinstance(schema, dict) else None
    if isinstance(title, str) and title.strip():
        return _single_line_text(title)
    return _single_line_text(key.replace("_", " ").strip().capitalize() or key)


def _is_scalar(value: Any) -> bool:
    return value is None or isinstance(value, (str, int, float, bool))


def _can_render_object_table(value: list[Any]) -> bool:
    for item in value:
        if not isinstance(item, dict):
            return False
        row = cast(dict[str, Any], item)
        if not all(_is_scalar(field) for field in row.values()):
            return False
    return True


def _table_block_for_objects(
    *, rows: list[dict[str, Any]], schema: dict[str, Any] | None
) -> _DocumentBlock:
    item_schema = _as_str_dict(schema.get("items") if schema is not None else None)
    item_properties = _as_str_dict(
        item_schema.get("properties") if item_schema is not None else None
    )
    columns = _ordered_table_columns(rows=rows, item_properties=item_properties)
    if not columns:
        return _DocumentBlock(kind="paragraph", text=_EMPTY_VALUE_PLACEHOLDER)

    header = tuple(
        _schema_label(
            column,
            _as_str_dict(item_properties.get(column) if item_properties else None),
        )
        for column in columns
    )
    body = tuple(
        tuple(_scalar_text(row.get(column)) for column in columns) for row in rows
    )
    return _DocumentBlock(kind="table", rows=(header, *body))


def _ordered_table_columns(
    *,
    rows: list[dict[str, Any]],
    item_properties: dict[str, Any] | None,
) -> list[str]:
    columns: list[str] = []
    if item_properties is not None:
        columns.extend(
            key
            for key in item_properties
            if any(key in row for row in rows) and key not in columns
        )
    for row in rows:
        columns.extend(key for key in row if key not in columns)
    return columns


def _scalar_text(value: Any) -> str:
    if value is None:
        return _EMPTY_VALUE_PLACEHOLDER
    if isinstance(value, bool):
        return "true" if value else "false"
    return _plain_text(str(value)) or _EMPTY_VALUE_PLACEHOLDER


def _plain_text(value: str) -> str:
    return value.replace("```", "'''").replace("\r\n", "\n").replace("\r", "\n").strip()


def _single_line_text(value: str) -> str:
    return (
        value.replace("```", "'''")
        .replace("\r\n", "\n")
        .replace("\r", "\n")
        .replace("\n", " ")
        .strip()
    )


def _as_str_dict(value: object) -> dict[str, Any] | None:
    return cast(dict[str, Any], value) if isinstance(value, dict) else None


def _trim_trailing_empty_blocks(blocks: list[_DocumentBlock]) -> list[_DocumentBlock]:
    trimmed = list(blocks)
    while trimmed and trimmed[-1].kind == "empty":
        trimmed.pop()
    return trimmed


def _render_docx(
    blocks: list[_DocumentBlock], *, step_order: int
) -> tuple[bytes, str, str]:
    import io

    from docx import Document

    # Use a repo-owned template so DOCX rendering does not depend on how
    # python-docx packaged its default template in the active environment.
    doc = (
        Document(str(_DOCX_TEMPLATE_PATH))
        if _DOCX_TEMPLATE_PATH.exists()
        else Document()
    )
    if not blocks:
        doc.add_paragraph("")
    else:
        _write_docx_blocks(doc=doc, blocks=blocks)
    buf = io.BytesIO()
    doc.save(buf)
    return (
        buf.getvalue(),
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        f"step_{step_order}_output.docx",
    )


def _write_docx_blocks(*, doc: Any, blocks: list[_DocumentBlock]) -> None:
    for block in blocks:
        _append_docx_block(doc=doc, block=block)


def _parse_markdown_blocks(lines: list[str]) -> list[_DocumentBlock]:
    if not lines:
        return [_DocumentBlock(kind="paragraph", text="")]

    blocks: list[_DocumentBlock] = []
    index = 0

    while index < len(lines):
        line = lines[index]
        stripped = line.strip()

        if stripped.startswith("```"):
            index += 1
            code_lines: list[str] = []
            while index < len(lines) and not lines[index].strip().startswith("```"):
                code_lines.append(lines[index].rstrip("\n"))
                index += 1
            if index < len(lines):
                index += 1
            blocks.append(_DocumentBlock(kind="code", text="\n".join(code_lines)))
            continue

        table_block, next_index = _consume_markdown_table_block(
            lines=lines,
            start_index=index,
        )
        if table_block is not None and next_index is not None:
            blocks.append(table_block)
            index = next_index
            continue

        if stripped == "":
            blocks.append(_DocumentBlock(kind="empty"))
            index += 1
            continue

        heading_match = _HEADING_PATTERN.match(stripped)
        if heading_match:
            level = min(len(heading_match.group(1)), 4)
            blocks.append(
                _DocumentBlock(
                    kind="heading",
                    text=heading_match.group(2).strip(),
                    level=level,
                )
            )
            index += 1
            continue

        bullet_match = _BULLET_PATTERN.match(stripped)
        if bullet_match:
            items: list[str] = []
            while index < len(lines):
                match = _BULLET_PATTERN.match(lines[index].strip())
                if match is None:
                    break
                items.append(match.group(1).strip())
                index += 1
            blocks.append(_DocumentBlock(kind="bullet_list", items=tuple(items)))
            continue

        numbered_match = _NUMBERED_PATTERN.match(stripped)
        if numbered_match:
            items: list[str] = []
            while index < len(lines):
                match = _NUMBERED_PATTERN.match(lines[index].strip())
                if match is None:
                    break
                items.append(match.group(1).strip())
                index += 1
            blocks.append(_DocumentBlock(kind="numbered_list", items=tuple(items)))
            continue

        paragraph_lines = [line]
        index += 1
        while index < len(lines):
            candidate = lines[index]
            candidate_stripped = candidate.strip()
            if candidate_stripped == "":
                break
            if candidate_stripped.startswith("```"):
                break
            if _HEADING_PATTERN.match(candidate_stripped):
                break
            if _BULLET_PATTERN.match(candidate_stripped):
                break
            if _NUMBERED_PATTERN.match(candidate_stripped):
                break
            if _is_markdown_table_start(lines=lines, start_index=index):
                break
            paragraph_lines.append(candidate)
            index += 1
        blocks.append(
            _DocumentBlock(kind="paragraph", text="\n".join(paragraph_lines).strip())
        )

    return blocks


def _consume_markdown_table_block(
    *,
    lines: list[str],
    start_index: int,
) -> tuple[_DocumentBlock | None, int | None]:
    if not _is_markdown_table_start(lines=lines, start_index=start_index):
        return None, None

    header_cells = _parse_markdown_table_row(lines[start_index])
    if not header_cells:
        return None, None

    rows: list[list[str]] = [header_cells]
    index = start_index + 2  # Skip separator row.

    while index < len(lines):
        line = lines[index]
        stripped = line.strip()
        if stripped == "" or "|" not in line:
            break
        cells = _parse_markdown_table_row(line)
        if not cells:
            break
        rows.append(cells)
        index += 1

    return (
        _DocumentBlock(
            kind="table",
            rows=tuple(tuple(cell for cell in row) for row in rows),
        ),
        index,
    )


def _append_docx_block(*, doc: Any, block: _DocumentBlock) -> None:
    if block.kind == "empty":
        doc.add_paragraph("")
        return
    if block.kind == "heading":
        doc.add_heading(block.text, level=block.level or 1)
        return
    if block.kind == "bullet_list":
        for item in block.items:
            doc.add_paragraph(item, style="List Bullet")
        return
    if block.kind == "numbered_list":
        for item in block.items:
            doc.add_paragraph(item, style="List Number")
        return
    if block.kind == "code":
        _append_code_block(doc=doc, code_lines=block.text.splitlines())
        return
    if block.kind == "table":
        _append_docx_table(doc=doc, rows=block.rows)
        return
    doc.add_paragraph(block.text)


def _append_docx_table(*, doc: Any, rows: tuple[tuple[str, ...], ...]) -> None:
    if not rows:
        return
    max_columns = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=max_columns)
    for row_idx, row in enumerate(rows):
        for col_idx in range(max_columns):
            value = row[col_idx] if col_idx < len(row) else ""
            table.cell(row_idx, col_idx).text = value


def _is_markdown_table_start(*, lines: list[str], start_index: int) -> bool:
    if start_index + 1 >= len(lines):
        return False
    header = lines[start_index]
    separator = lines[start_index + 1]
    return (
        "|" in header and _MARKDOWN_TABLE_SEPARATOR.match(separator.strip()) is not None
    )


def _parse_markdown_table_row(line: str) -> list[str]:
    stripped = line.strip()
    if not stripped:
        return []
    if stripped.startswith("|"):
        stripped = stripped[1:]
    if stripped.endswith("|"):
        stripped = stripped[:-1]
    escaped_pipe = "\0ENEO_ESCAPED_PIPE\0"
    cells = stripped.replace("\\|", escaped_pipe).split("|")
    return [cell.replace(escaped_pipe, "|").strip() for cell in cells]


def _append_code_block(*, doc: Any, code_lines: list[str]) -> None:
    paragraph = doc.add_paragraph("\n".join(code_lines))
    if not paragraph.runs:
        return
    for run in paragraph.runs:
        run.font.name = "Courier New"


def _write_pdf_blocks(
    *,
    pdf: Any,
    blocks: list[_DocumentBlock],
    font_family: str,
    uses_core_font: bool,
) -> None:
    for block in blocks:
        _append_pdf_block(
            pdf=pdf,
            block=block,
            font_family=font_family,
            uses_core_font=uses_core_font,
        )


def _append_pdf_block(
    *,
    pdf: Any,
    block: _DocumentBlock,
    font_family: str,
    uses_core_font: bool,
) -> None:
    if block.kind == "empty":
        pdf.ln(3)
        return
    if block.kind == "heading":
        size_by_level = {1: 16, 2: 14, 3: 12, 4: 11}
        size = size_by_level.get(block.level, 11)
        pdf.set_font(font_family, size=size)
        _pdf_multi_cell(
            pdf,
            height=max(7, int(size * 0.55)),
            text=_pdf_text(block.text, uses_core_font),
        )
        pdf.ln(1)
        return
    if block.kind == "bullet_list":
        pdf.set_font(font_family, size=11)
        for item in block.items:
            _pdf_multi_cell(
                pdf,
                height=6,
                text=_pdf_text(f"- {item}", uses_core_font),
            )
        pdf.ln(1)
        return
    if block.kind == "numbered_list":
        pdf.set_font(font_family, size=11)
        for index, item in enumerate(block.items, start=1):
            _pdf_multi_cell(
                pdf,
                height=6,
                text=_pdf_text(f"{index}. {item}", uses_core_font),
            )
        pdf.ln(1)
        return
    if block.kind == "code":
        pdf.set_font(font_family, size=9)
        for line in block.text.splitlines() or [""]:
            _pdf_multi_cell(pdf, height=5, text=_pdf_text(line, uses_core_font))
        pdf.ln(1)
        return
    if block.kind == "table":
        _append_pdf_table(
            pdf=pdf,
            rows=block.rows,
            font_family=font_family,
            uses_core_font=uses_core_font,
        )
        return
    pdf.set_font(font_family, size=11)
    _pdf_multi_cell(pdf, height=6, text=_pdf_text(block.text, uses_core_font))
    pdf.ln(1)


def _append_pdf_table(
    *,
    pdf: Any,
    rows: tuple[tuple[str, ...], ...],
    font_family: str,
    uses_core_font: bool,
) -> None:
    if not rows:
        return
    max_columns = max(len(row) for row in rows)
    usable_width = pdf.w - pdf.l_margin - pdf.r_margin
    pdf.set_font(font_family, size=9)
    column_widths = _pdf_table_column_widths(
        pdf=pdf,
        rows=rows,
        usable_width=usable_width,
        max_columns=max_columns,
    )
    line_height = 5
    header_layout: tuple[list[list[str]], float] | None = None
    if len(rows) > 1:
        header_layout = _pdf_table_row_layout(
            pdf=pdf,
            row=rows[0],
            column_widths=column_widths,
            max_columns=max_columns,
            line_height=line_height,
            uses_core_font=uses_core_font,
        )

    for row_index, row in enumerate(rows):
        pdf.set_font(font_family, size=9)
        wrapped_cells, row_height = _pdf_table_row_layout(
            pdf=pdf,
            row=row,
            column_widths=column_widths,
            max_columns=max_columns,
            line_height=line_height,
            uses_core_font=uses_core_font,
        )
        repeated_header_height = (
            header_layout[1] if row_index > 0 and header_layout is not None else 0
        )
        started_new_page = _ensure_pdf_row_space(
            pdf,
            row_height=row_height + repeated_header_height,
        )
        if started_new_page and row_index > 0 and header_layout is not None:
            header_cells, header_height = header_layout
            _draw_pdf_table_row(
                pdf=pdf,
                wrapped_cells=header_cells,
                row_height=header_height,
                column_widths=column_widths,
                line_height=line_height,
            )
        _draw_pdf_table_row(
            pdf=pdf,
            wrapped_cells=wrapped_cells,
            row_height=row_height,
            column_widths=column_widths,
            line_height=line_height,
        )
    pdf.ln(2)


def _pdf_table_column_widths(
    *,
    pdf: Any,
    rows: tuple[tuple[str, ...], ...],
    usable_width: float,
    max_columns: int,
) -> list[float]:
    if max_columns <= 0:
        return []
    equal_width = usable_width / max_columns
    min_width = min(28.0, equal_width)
    max_width = usable_width if max_columns == 1 else usable_width * 0.55
    desired_widths: list[float] = []
    for column_index in range(max_columns):
        widest_value = 0.0
        for row in rows:
            value = row[column_index] if column_index < len(row) else ""
            for line in value.splitlines() or [""]:
                widest_value = max(widest_value, float(pdf.get_string_width(line)) + 4)
        desired_widths.append(min(max(widest_value, min_width), max_width))

    desired_total = sum(desired_widths)
    if desired_total <= 0:
        return [equal_width] * max_columns
    if desired_total <= usable_width:
        extra_width = usable_width - desired_total
        return [
            width + (extra_width * (width / desired_total)) for width in desired_widths
        ]

    minimum_total = min_width * max_columns
    if minimum_total >= usable_width:
        return [equal_width] * max_columns

    flexible_total = sum(max(width - min_width, 0.0) for width in desired_widths)
    if flexible_total <= 0:
        return [equal_width] * max_columns

    overflow = desired_total - usable_width
    return [
        width - (overflow * (max(width - min_width, 0.0) / flexible_total))
        for width in desired_widths
    ]


def _pdf_table_row_layout(
    *,
    pdf: Any,
    row: tuple[str, ...],
    column_widths: list[float],
    max_columns: int,
    line_height: int,
    uses_core_font: bool,
) -> tuple[list[list[str]], float]:
    values = [
        row[col_idx] if col_idx < len(row) else "" for col_idx in range(max_columns)
    ]
    wrapped_cells = [
        _wrap_pdf_cell_text(
            pdf,
            text=_pdf_text(value, uses_core_font),
            max_width=column_widths[col_idx] - 2,
        )
        for col_idx, value in enumerate(values)
    ]
    row_height = max(len(lines) for lines in wrapped_cells) * line_height + 2
    return wrapped_cells, row_height


def _draw_pdf_table_row(
    *,
    pdf: Any,
    wrapped_cells: list[list[str]],
    row_height: float,
    column_widths: list[float],
    line_height: int,
) -> None:
    row_x = pdf.l_margin
    row_y = pdf.get_y()
    offset_x = row_x
    for col_idx, lines in enumerate(wrapped_cells):
        cell_width = column_widths[col_idx]
        pdf.rect(offset_x, row_y, cell_width, row_height)
        pdf.set_xy(offset_x + 1, row_y + 1)
        pdf.multi_cell(cell_width - 2, line_height, "\n".join(lines), border=0)
        offset_x += cell_width
    pdf.set_xy(row_x, row_y + row_height)


def _pdf_text(text: str, uses_core_font: bool) -> str:
    return _text_for_core_pdf_font(text) if uses_core_font else text


def _pdf_multi_cell(pdf: Any, *, height: int, text: str) -> None:
    pdf.set_x(pdf.l_margin)
    pdf.multi_cell(0, height, text)
    pdf.set_x(pdf.l_margin)


def _ensure_pdf_row_space(pdf: Any, *, row_height: float) -> bool:
    page_break_trigger = getattr(pdf, "page_break_trigger", None)
    if page_break_trigger is not None and pdf.get_y() + row_height > page_break_trigger:
        pdf.add_page()
        pdf.set_x(pdf.l_margin)
        return True
    return False


def _wrap_pdf_cell_text(pdf: Any, *, text: str, max_width: float) -> list[str]:
    if max_width <= 0:
        return [text]

    lines: list[str] = []
    for source_line in text.splitlines() or [""]:
        words = source_line.split()
        if not words:
            lines.append("")
            continue

        current = ""
        for word in words:
            candidate = word if not current else f"{current} {word}"
            if pdf.get_string_width(candidate) <= max_width:
                current = candidate
                continue
            if current:
                lines.append(current)
            if pdf.get_string_width(word) <= max_width:
                current = word
            else:
                broken_word = _break_pdf_word(pdf, word=word, max_width=max_width)
                lines.extend(broken_word[:-1])
                current = broken_word[-1] if broken_word else ""
        if current:
            lines.append(current)
    return lines or [""]


def _break_pdf_word(pdf: Any, *, word: str, max_width: float) -> list[str]:
    parts: list[str] = []
    current = ""
    for char in word:
        candidate = f"{current}{char}"
        if current and pdf.get_string_width(candidate) > max_width:
            parts.append(current)
            current = char
        else:
            current = candidate
    if current:
        parts.append(current)
    return parts or [word]
