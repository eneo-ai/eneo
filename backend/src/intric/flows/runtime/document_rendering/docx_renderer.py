from __future__ import annotations

from pathlib import Path
from typing import Any, Sequence

from intric.flows.runtime.document_rendering.blocks import (
    DocumentBlock,
    InlineRuns,
    InlineTextRun,
)
from intric.flows.runtime.document_rendering.renderers import RenderedDocument

_DOCX_TEMPLATE_PATH = (
    Path(__file__).resolve().parent.parent / "templates" / "default.docx"
)
_DOCX_MIMETYPE = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)


class DocxDocumentRenderer:
    output_type = "docx"

    def __init__(self, *, template_path: Path = _DOCX_TEMPLATE_PATH) -> None:
        self._template_path = template_path

    def render(
        self,
        blocks: Sequence[DocumentBlock],
        *,
        step_order: int,
    ) -> RenderedDocument:
        import io

        from docx import Document

        doc = (
            Document(str(self._template_path))
            if self._template_path.exists()
            else Document()
        )
        if not blocks:
            doc.add_paragraph("")
        else:
            for block in blocks:
                _append_docx_block(doc=doc, block=block)
        buf = io.BytesIO()
        doc.save(buf)
        return RenderedDocument(
            blob=buf.getvalue(),
            mimetype=_DOCX_MIMETYPE,
            filename=f"step_{step_order}_output.docx",
        )


def _append_docx_block(*, doc: Any, block: DocumentBlock) -> None:
    if block.kind == "empty":
        doc.add_paragraph("")
        return
    if block.kind == "heading":
        paragraph = doc.add_heading("", level=block.level or 1)
        _append_runs_to_paragraph(
            paragraph=paragraph,
            runs=block.runs,
            fallback=block.text,
        )
        return
    if block.kind == "bullet_list":
        for index, item in enumerate(block.items):
            paragraph = doc.add_paragraph(style="List Bullet")
            _append_runs_to_paragraph(
                paragraph=paragraph,
                runs=_runs_at(block.item_runs, index),
                fallback=item,
            )
        return
    if block.kind == "numbered_list":
        for index, item in enumerate(block.items):
            paragraph = doc.add_paragraph(style="List Number")
            _append_runs_to_paragraph(
                paragraph=paragraph,
                runs=_runs_at(block.item_runs, index),
                fallback=item,
            )
        return
    if block.kind == "code":
        _append_code_block(doc=doc, code_lines=block.text.splitlines())
        return
    if block.kind == "table":
        _append_docx_table(doc=doc, rows=block.rows, row_runs=block.row_runs)
        return
    paragraph = doc.add_paragraph()
    _append_runs_to_paragraph(
        paragraph=paragraph,
        runs=block.runs,
        fallback=block.text,
    )


def _append_docx_table(
    *,
    doc: Any,
    rows: tuple[tuple[str, ...], ...],
    row_runs: tuple[tuple[InlineRuns, ...], ...],
) -> None:
    if not rows:
        return
    max_columns = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=max_columns)
    for row_idx, row in enumerate(rows):
        for col_idx in range(max_columns):
            value = row[col_idx] if col_idx < len(row) else ""
            paragraph = table.cell(row_idx, col_idx).paragraphs[0]
            _append_runs_to_paragraph(
                paragraph=paragraph,
                runs=_cell_runs_at(row_runs, row_idx, col_idx),
                fallback=value,
                force_bold=row_idx == 0,
            )


def _append_code_block(*, doc: Any, code_lines: list[str]) -> None:
    paragraph = doc.add_paragraph("\n".join(code_lines))
    if not paragraph.runs:
        return
    for run in paragraph.runs:
        run.font.name = "Courier New"


def _append_runs_to_paragraph(
    *,
    paragraph: Any,
    runs: InlineRuns,
    fallback: str,
    force_bold: bool = False,
) -> None:
    if not runs:
        if fallback:
            run = paragraph.add_run(fallback)
            run.bold = force_bold
        return
    for inline_run in runs:
        _append_run(
            paragraph=paragraph,
            inline_run=inline_run,
            force_bold=force_bold,
        )


def _append_run(
    *,
    paragraph: Any,
    inline_run: InlineTextRun,
    force_bold: bool,
) -> None:
    run = paragraph.add_run(inline_run.text)
    run.bold = inline_run.bold or force_bold
    run.italic = inline_run.italic
    run.font.strike = inline_run.strikethrough
    if inline_run.code:
        run.font.name = "Courier New"


def _runs_at(runs: tuple[InlineRuns, ...], index: int) -> InlineRuns:
    return runs[index] if index < len(runs) else ()


def _cell_runs_at(
    rows: tuple[tuple[InlineRuns, ...], ...],
    row_index: int,
    column_index: int,
) -> InlineRuns:
    if row_index >= len(rows):
        return ()
    row = rows[row_index]
    return row[column_index] if column_index < len(row) else ()
