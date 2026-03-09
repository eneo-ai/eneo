from __future__ import annotations

import io
import logging
import re
import tempfile
from pathlib import Path
from typing import Any

from docx import Document
from docx.oxml.ns import qn
from jinja2.sandbox import SandboxedEnvironment

from intric.files.docx_template_validation import validate_docx_template_archive
from intric.main.exceptions import TypedIOValidationException

logger = logging.getLogger(__name__)

_PLACEHOLDER_PATTERN = re.compile(r"\{\{\s*([^{}]+)\s*\}\}")
_SAFE_TEMPLATE_NAME_PATTERN = re.compile(r"^[A-Za-z_][A-Za-z0-9_]*(\.[A-Za-z_][A-Za-z0-9_]*)*$")
_DOCX_MIMETYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

# OOXML header/footer reference types → python-docx section properties
_HEADER_TYPE_TO_ATTR = {
    "default": "header",
    "first": "first_page_header",
    "even": "even_page_header",
}
_FOOTER_TYPE_TO_ATTR = {
    "default": "footer",
    "first": "first_page_footer",
    "even": "even_page_footer",
}


def _iter_section_headers(section: Any) -> list[Any]:
    """Return headers explicitly defined in the section XML.

    Only accesses the python-docx property matching the reference type
    (default/first/even) to avoid triggering implicit header creation
    which requires template files that may not be installed.
    """
    headers = []
    for ref in section._sectPr.findall(qn("w:headerReference")):
        hdr_type = ref.get(qn("w:type"), "default")
        attr = _HEADER_TYPE_TO_ATTR.get(hdr_type)
        if attr and hasattr(section, attr):
            try:
                headers.append(getattr(section, attr))
            except (FileNotFoundError, KeyError):
                logger.warning("Skipping inaccessible %s header in section", hdr_type)
    return headers


def _iter_section_footers(section: Any) -> list[Any]:
    """Return footers explicitly defined in the section XML.

    Only accesses the python-docx property matching the reference type
    (default/first/even) to avoid triggering implicit footer creation
    which requires template files that may not be installed.
    """
    footers = []
    for ref in section._sectPr.findall(qn("w:footerReference")):
        ftr_type = ref.get(qn("w:type"), "default")
        attr = _FOOTER_TYPE_TO_ATTR.get(ftr_type)
        if attr and hasattr(section, attr):
            try:
                footers.append(getattr(section, attr))
            except (FileNotFoundError, KeyError):
                logger.warning("Skipping inaccessible %s footer in section", ftr_type)
    return footers


def inspect_docx_template_bytes(
    template_bytes: bytes,
    *,
    filename: str,
) -> list[dict[str, str | None]]:
    validate_docx_template_archive(template_bytes, filename=filename)
    document = Document(io.BytesIO(template_bytes))

    discovered: list[dict[str, str | None]] = []
    seen: set[tuple[str, str]] = set()

    def record(text: str, *, location: str) -> None:
        preview = text.strip()[:160] or None
        for match in _PLACEHOLDER_PATTERN.finditer(text):
            name = match.group(1).strip()
            key = (name, location)
            if not name or key in seen:
                continue
            seen.add(key)
            discovered.append({"name": name, "location": location, "preview": preview})

    for paragraph in document.paragraphs:
        record(paragraph.text, location="body")
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                record(cell.text, location="table")
    for section in document.sections:
        for hdr in _iter_section_headers(section):
            for paragraph in hdr.paragraphs:
                record(paragraph.text, location="header")
            for table in hdr.tables:
                for row in table.rows:
                    for cell in row.cells:
                        record(cell.text, location="header")
        for ftr in _iter_section_footers(section):
            for paragraph in ftr.paragraphs:
                record(paragraph.text, location="footer")
            for table in ftr.tables:
                for row in table.rows:
                    for cell in row.cells:
                        record(cell.text, location="footer")

    return discovered


def render_docx_template(
    *,
    template_bytes: bytes,
    context: dict[str, Any],
    step_order: int,
) -> tuple[bytes, str, str]:
    placeholders = inspect_docx_template_bytes(template_bytes, filename=f"step_{step_order}.docx")
    required_names = {str(item["name"]) for item in placeholders}
    missing = sorted(name for name in required_names if name not in context)
    if missing:
        raise TypedIOValidationException(
            f"Unresolved template placeholders: {', '.join(missing)}",
            code="typed_io_template_render_failed",
        )

    try:
        from docxtpl import DocxTemplate
    except Exception as exc:  # pragma: no cover - environment guard
        raise TypedIOValidationException(
            f"DOCX template rendering dependency is unavailable: {exc}",
            code="typed_io_template_render_failed",
        ) from exc

    normalized_template_bytes, alias_by_name = _normalize_docx_template_placeholders(
        template_bytes,
        required_names=required_names,
    )
    render_context = _build_docxtpl_context(context)
    for name, alias in alias_by_name.items():
        render_context[alias] = context[name]

    with tempfile.TemporaryDirectory(prefix="flow-docx-template-") as temp_dir:
        template_path = Path(temp_dir) / "template.docx"
        output_path = Path(temp_dir) / f"step_{step_order}_output.docx"
        template_path.write_bytes(normalized_template_bytes)
        try:
            template = DocxTemplate(str(template_path))
            template.render(render_context, jinja_env=SandboxedEnvironment(autoescape=False))
            template.save(str(output_path))
            blob = output_path.read_bytes()
        except TypedIOValidationException:
            raise
        except Exception as exc:
            message = str(exc)
            if "UndefinedError" in message or "is undefined" in message:
                raise TypedIOValidationException(
                    "Template rendering failed. Check that placeholders are written directly in the DOCX without extra formatting inside the braces.",
                    code="typed_io_template_render_failed",
                ) from exc
            raise TypedIOValidationException(
                f"DOCX template render failed: {exc}",
                code="typed_io_template_render_failed",
            ) from exc

    return blob, _DOCX_MIMETYPE, f"step_{step_order}_output.docx"


def _normalize_docx_template_placeholders(
    template_bytes: bytes,
    *,
    required_names: set[str],
) -> tuple[bytes, dict[str, str]]:
    alias_by_name = {
        name: f"placeholder_{index}"
        for index, name in enumerate(sorted(required_names), start=1)
        if _SAFE_TEMPLATE_NAME_PATTERN.fullmatch(name) is None
    }
    if not alias_by_name:
        return template_bytes, {}

    document = Document(io.BytesIO(template_bytes))
    for paragraph in _iter_story_paragraphs(document):
        _replace_placeholder_aliases_in_paragraph(paragraph, alias_by_name)

    output = io.BytesIO()
    document.save(output)
    return output.getvalue(), alias_by_name


def _build_docxtpl_context(context: dict[str, Any]) -> dict[str, Any]:
    """Expose flat placeholder bindings both directly and as nested dotted aliases."""
    render_context: dict[str, Any] = dict(context)

    for key, value in context.items():
        if "." not in key:
            continue

        path = [segment.strip() for segment in key.split(".") if segment.strip()]
        if len(path) < 2:
            continue

        current: dict[str, Any] = render_context
        for segment in path[:-1]:
            existing = current.get(segment)
            if existing is None:
                next_level: dict[str, Any] = {}
                current[segment] = next_level
                current = next_level
                continue
            if not isinstance(existing, dict):
                current = {}
                break
            current = existing
        else:
            current[path[-1]] = value

    return render_context


def _iter_story_paragraphs(container: Any):
    for paragraph in getattr(container, "paragraphs", []):
        yield paragraph
    for table in getattr(container, "tables", []):
        for row in table.rows:
            for cell in row.cells:
                yield from _iter_story_paragraphs(cell)
    for section in getattr(container, "sections", []):
        for hdr in _iter_section_headers(section):
            yield from _iter_story_paragraphs(hdr)
        for ftr in _iter_section_footers(section):
            yield from _iter_story_paragraphs(ftr)


def _replace_placeholder_aliases_in_paragraph(
    paragraph: Any,
    alias_by_name: dict[str, str],
) -> None:
    original_text = paragraph.text
    if not original_text or "{{" not in original_text:
        return

    replaced_text = _PLACEHOLDER_PATTERN.sub(
        lambda match: "{{" + alias_by_name.get(match.group(1).strip(), match.group(1).strip()) + "}}",
        original_text,
    )
    if replaced_text == original_text:
        return

    text_nodes = list(paragraph._p.iter(qn("w:t")))
    if not text_nodes:
        return

    text_nodes[0].text = replaced_text
    for node in text_nodes[1:]:
        node.text = ""


def extract_docx_text(document_bytes: bytes) -> str:
    document = Document(io.BytesIO(document_bytes))
    segments: list[str] = []
    segments.extend(
        paragraph.text.strip()
        for paragraph in document.paragraphs
        if paragraph.text.strip()
    )
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    segments.append(text)
    for section in document.sections:
        for hdr in _iter_section_headers(section):
            for paragraph in hdr.paragraphs:
                text = paragraph.text.strip()
                if text:
                    segments.append(text)
        for ftr in _iter_section_footers(section):
            for paragraph in ftr.paragraphs:
                text = paragraph.text.strip()
                if text:
                    segments.append(text)
    return "\n\n".join(segments)


def extract_docx_template_text_preview(template_bytes: bytes) -> str:
    return extract_docx_text(template_bytes)[:2000]
