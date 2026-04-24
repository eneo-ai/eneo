"""Deterministic signal extractor for AI Builder attachment observations.

The observation pipeline asks this module, before any LLM call, what
the raw bytes alone can tell us: the canonical MIME, a filename
extension, the size in bytes, and — for text-like payloads — a
bullet-line density that hints at whether the content is dense prose
or a list. Per-mime extractors for CSV / media plug into the dispatcher
below; text, markdown, DOCX, PDF, and XLSX are handled here.

Unsupported or legacy MIMEs still produce a ``DeterministicSignals``
with the file-level fields populated and per-mime fields at their
defaults, so upstream callers never need a ``None`` check on the
top-level return. Rejection of legacy uploads happens at the upload
boundary via ``files.mime_support.classify_mime``; here we stay
permissive so the cache key remains derivable for every byte stream
that reaches the pipeline.
"""

from __future__ import annotations

import datetime as dt
import re
import zipfile
from io import BytesIO

import pdfplumber
from docx import Document
from docx.opc.exceptions import PackageNotFoundError
from openpyxl import load_workbook

from intric.files.extensions import MIMETYPE_EXTENSIONS_MAPPER
from intric.files.mime_support import canonicalize_mime
from intric.flows.ai_builder.attachment_observation import (
    DeterministicSignals,
    Heading,
    TableDimension,
)

_TEXT_MIMES: frozenset[str] = frozenset({"text/plain", "text/markdown"})
_DOCX_MIME: str = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)
_PDF_MIME: str = "application/pdf"
_XLSX_MIME: str = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
_BULLET_LINE = re.compile(r"^\s*[-*+•·]\s+\S")
# Jinja-style `{{ var }}` placeholder tokens. The surrounding pattern
# is deliberately permissive (identifier + optional whitespace) so the
# set surfaces what template-style documents typically expose, not a
# full Jinja parse.
_PLACEHOLDER_TOKEN = re.compile(r"\{\{\s*([A-Za-z_][A-Za-z0-9_]*)\s*\}\}")


def extract_deterministic_signals(
    *,
    raw_bytes: bytes,
    mime: str,
    filename: str,
) -> DeterministicSignals:
    canonical_mime = canonicalize_mime(mime) or (mime or "")
    extension = _infer_extension(filename=filename, mime=canonical_mime)
    size_bytes = len(raw_bytes)

    if canonical_mime in _TEXT_MIMES:
        return DeterministicSignals(
            mime_type=canonical_mime,
            extension=extension,
            size_bytes=size_bytes,
            bullet_density=_bullet_density(raw_bytes),
        )

    if canonical_mime == _DOCX_MIME:
        return _docx_signals(
            raw_bytes=raw_bytes,
            canonical_mime=canonical_mime,
            extension=extension,
            size_bytes=size_bytes,
        )

    if canonical_mime == _PDF_MIME:
        return _pdf_signals(
            raw_bytes=raw_bytes,
            canonical_mime=canonical_mime,
            extension=extension,
            size_bytes=size_bytes,
        )

    if canonical_mime == _XLSX_MIME:
        return _xlsx_signals(
            raw_bytes=raw_bytes,
            canonical_mime=canonical_mime,
            extension=extension,
            size_bytes=size_bytes,
        )

    return DeterministicSignals(
        mime_type=canonical_mime,
        extension=extension,
        size_bytes=size_bytes,
    )


def _docx_signals(
    *,
    raw_bytes: bytes,
    canonical_mime: str,
    extension: str,
    size_bytes: int,
) -> DeterministicSignals:
    """Extract DOCX-specific signals without rendering the document.

    Headings come from paragraph styles (level 1..6); placeholder
    tokens are collected across every paragraph and every table cell
    so a table-embedded `{{var}}` is not missed; table dimensions are
    read from the document's table list. Page count is intentionally
    not computed — DOCX doesn't carry an authoritative page count
    until rendered.
    """
    try:
        document = Document(BytesIO(raw_bytes))
    except (PackageNotFoundError, zipfile.BadZipFile, KeyError, ValueError):
        # Malformed bytes — the dispatcher must not crash. Fall back
        # to file-level-only signals so the pipeline still has a
        # cache-keyable return. Per-mime fields stay at their
        # defaults (None / empty list).
        return DeterministicSignals(
            mime_type=canonical_mime,
            extension=extension,
            size_bytes=size_bytes,
        )

    headings: list[Heading] = []
    placeholder_tokens: list[str] = []
    seen_tokens: set[str] = set()

    def _scan_text(text: str) -> None:
        for match in _PLACEHOLDER_TOKEN.finditer(text):
            token = match.group(1)
            if token not in seen_tokens:
                seen_tokens.add(token)
                placeholder_tokens.append(token)

    for paragraph in document.paragraphs:
        style_name = paragraph.style.name if paragraph.style is not None else ""
        level = _heading_level_from_style(style_name or "")
        if level is not None and paragraph.text.strip():
            headings.append(Heading(level=level, text=paragraph.text.strip()))
        _scan_text(paragraph.text)

    table_dimensions: list[TableDimension] = []
    for table in document.tables:
        rows = len(table.rows)
        cols = len(table.columns)
        table_dimensions.append(TableDimension(rows=rows, cols=cols))
        for row in table.rows:
            for cell in row.cells:
                _scan_text(cell.text)

    return DeterministicSignals(
        mime_type=canonical_mime,
        extension=extension,
        size_bytes=size_bytes,
        heading_tree=headings,
        table_count=len(table_dimensions),
        table_dimensions=table_dimensions,
        placeholder_tokens=placeholder_tokens,
    )


def _pdf_signals(
    *,
    raw_bytes: bytes,
    canonical_mime: str,
    extension: str,
    size_bytes: int,
) -> DeterministicSignals:
    """Extract PDF-specific signals without OCR.

    Page count is authoritative from the page list. `is_scanned_pdf`
    is a proxy — True when no page returns any extractable text, which
    covers image-only scans and genuinely empty PDFs alike. The
    planner uses this flag to avoid asking a scanned attachment to
    populate a text-summary archetype without a user acknowledgement;
    distinguishing a true scan from a blank PDF requires OCR and is
    explicitly out of scope here.
    """
    try:
        with pdfplumber.open(BytesIO(raw_bytes)) as pdf:
            page_count = len(pdf.pages)
            any_text = any((page.extract_text() or "").strip() for page in pdf.pages)
    except Exception:
        # pdfplumber + pdfminer raise a wide variety of exceptions on
        # malformed bytes (PdfminerException, PDFSyntaxError,
        # ValueError, AssertionError from the underlying parser).
        # Catching the umbrella here matches the upstream-permissive
        # MIME classification contract: the dispatcher must never
        # crash on bytes that reached it, and downstream code is
        # cache-keyed on file-level signals alone when per-mime
        # parsing fails.
        return DeterministicSignals(
            mime_type=canonical_mime,
            extension=extension,
            size_bytes=size_bytes,
        )

    return DeterministicSignals(
        mime_type=canonical_mime,
        extension=extension,
        size_bytes=size_bytes,
        page_count=page_count,
        is_scanned_pdf=not any_text,
    )


def _xlsx_signals(
    *,
    raw_bytes: bytes,
    canonical_mime: str,
    extension: str,
    size_bytes: int,
) -> DeterministicSignals:
    """Extract XLSX headers, per-column types, and data row count.

    Headers are the first row's values, column types are inferred from
    the data rows, and `row_count` excludes the header. Only the
    active sheet is inspected — the multi-sheet case is deferred until
    a golden requires it, since conflating sheets into a single header
    list would misrepresent the document's structure to the planner.
    """
    try:
        workbook = load_workbook(BytesIO(raw_bytes), read_only=True, data_only=True)
    except (zipfile.BadZipFile, KeyError, ValueError, OSError):
        return DeterministicSignals(
            mime_type=canonical_mime,
            extension=extension,
            size_bytes=size_bytes,
        )

    try:
        sheet = workbook.active
        if sheet is None:
            return DeterministicSignals(
                mime_type=canonical_mime,
                extension=extension,
                size_bytes=size_bytes,
            )

        rows_iter = sheet.iter_rows(values_only=True)
        header_row = next(rows_iter, None)
        if header_row is None:
            return DeterministicSignals(
                mime_type=canonical_mime,
                extension=extension,
                size_bytes=size_bytes,
                spreadsheet_headers=[],
                spreadsheet_column_types=[],
                row_count=0,
            )

        headers = [str(cell) if cell is not None else "" for cell in header_row]
        column_values: list[list[object]] = [[] for _ in headers]
        row_count = 0
        for data_row in rows_iter:
            row_count += 1
            for index, value in enumerate(data_row):
                if index < len(column_values):
                    column_values[index].append(value)

        column_types = [_infer_column_type(values) for values in column_values]

        return DeterministicSignals(
            mime_type=canonical_mime,
            extension=extension,
            size_bytes=size_bytes,
            spreadsheet_headers=headers,
            spreadsheet_column_types=column_types,
            row_count=row_count,
        )
    finally:
        workbook.close()


def _infer_column_type(values: list[object]) -> str:
    """Classify a spreadsheet column by scanning its non-empty cell values.

    All-empty → `"empty"`. All values of a single inferred type →
    that type's name. Any mix of types → `"mixed"`, because a
    JSON-schema extraction flow cannot safely claim a single type for
    a column that contains divergent cell kinds.
    """
    seen: set[str] = set()
    for value in values:
        if value is None or (isinstance(value, str) and value == ""):
            continue
        # bool must be checked before int — `isinstance(True, int)` is True
        # in Python, which would misclassify booleans as numeric.
        if isinstance(value, bool):
            seen.add("boolean")
        elif isinstance(value, (int, float)):
            seen.add("numeric")
        elif isinstance(value, (dt.datetime, dt.date)):
            seen.add("date")
        elif isinstance(value, str):
            seen.add("text")
        else:
            seen.add("other")
    if not seen:
        return "empty"
    if len(seen) == 1:
        return next(iter(seen))
    return "mixed"


def _heading_level_from_style(style_name: str) -> int | None:
    """Map a docx paragraph style name to a heading level.

    `python-docx` exposes style names like `"Heading 1"`, `"Heading 2"`,
    ..., `"Heading 9"`. Anything outside 1..6 is ignored (the schema
    constrains `Heading.level` to that range, and level-7+ is a
    style-local extension, not a structural signal).
    """
    if not style_name.lower().startswith("heading "):
        return None
    try:
        level = int(style_name.rsplit(" ", 1)[-1])
    except ValueError:
        return None
    if 1 <= level <= 6:
        return level
    return None


def _infer_extension(*, filename: str, mime: str) -> str:
    suffix = _filename_suffix(filename)
    if suffix:
        return suffix
    extensions = MIMETYPE_EXTENSIONS_MAPPER.get(mime, [])
    if extensions:
        return extensions[0].lstrip(".")
    return ""


def _filename_suffix(filename: str) -> str:
    if "." not in filename:
        return ""
    # A hidden filename like ".gitignore" has no informative suffix —
    # the leading dot is a visibility marker, not an extension.
    if filename.startswith(".") and filename.count(".") == 1:
        return ""
    _, _, suffix = filename.rpartition(".")
    return suffix.lower()


def _bullet_density(raw_bytes: bytes) -> float | None:
    if not raw_bytes:
        return None
    text = raw_bytes.decode("utf-8", errors="replace")
    lines = [line for line in text.splitlines() if line.strip()]
    if not lines:
        return None
    bullet_count = sum(1 for line in lines if _BULLET_LINE.match(line))
    return bullet_count / len(lines)
